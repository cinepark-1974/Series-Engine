"""
👖 BLUE JEANS SERIES ENGINE v2.0.8 — main.py
시즌 아크 → 에피소드 씬 플랜 → 비트 집필 파이프라인
© 2026 BLUE JEANS PICTURES

★ v2.0 본 작업 — 완전한 엔진
  [패치 A] Writer Engine 양식 빌더 이식 (v1.9 패치본 유지)
  [패치 B] 회별 특수 모듈 — OPENING/MIDSEASON TWIST/LOWEST POINT/FINALE MASTERY
  [패치 C] INSERT 시스템 + PROP 연속성 + GENRE BOOSTER 9장르 + HELPER 룰
  [패치 D] 시즌 표현 누적 차단 시스템 (AI Escape A17 해결)
  + SCENE RULES 강화 (씬 분할 / 씬 번호 / 대사 형식)
  + 회별 사건/시즌 비중 룰 (EPISODE_FOCUS_RATIO)

★ v2.0.1 미세 패치
  [패치 E] 캐릭터+대사 분리 복구 강화 — 한국 시나리오 표준 양식 정합
  [패치 F] DOCX/TXT 파일명 규칙 갱신 — 각본_제목_v1.0_날짜_시간

★ v2.0.2 미세 패치
  [패치 G] 씬 헤딩 시간 표기 5단계 표준 강제
    · 5단계: 새벽 / 오전 / 오후 / 저녁 / 밤
    · 시·분 단위·세분화·"낮" 단독 금지

★ v2.0.3 미세 패치
  [패치 H] DOCX 출력 모드 토글 — 비트헤더 선택적 포함
    · 최종 모드 (기본): 비트 헤더 제거
    · 집필 모드: 비트 헤더 포함

★ v2.0.4 미세 패치
  [패치 I] 외부 모니터링 / 쇼러너 노트 입력
    · STEP 2 시즌 아크 영역에 expander 추가
    · build_season_arc_prompt에 monitoring_feedback / showrunner_notes 인자 전달
    · 시즌 아크 설계 시 외부 피드백을 본문에 자동 주입
    · 백업 JSON 스키마 확장 (2개 키 추가)
  [Hotfix 1] build_locked_block NameError 수정 (2026-05-17)
    · prompt.py L61~91에 잘못 박혀 있던 죽은 코드 제거
    · `p`, `core` 변수 미정의 NameError 해소
    · LOCKED 5종 확장 처리는 extract_from_creator_json_series에 정상 존재 (중복 코드였음)
    · 호출부 두 곳(_get_locked_block / rewrite 모드) 모두 정상 동작

★ v2.0.5 미세 패치
  [패치 J] 작품 제목 입력 + 백업 파일명/사이드바/PDF 푸터 자동 주입
  [패치 K] 리라이트 모드 자동 로더 (Creator JSON + Series JSON)

★ v2.0.6 미세 패치 (2026-05-20)
  [Hotfix 2] 백업 복원 시 monitoring_feedback / showrunner_notes
            "cannot be modified after the widget with key ... is instantiated"
            오류 해소
    · 원인: v2.0.4에서 추가된 두 신규 위젯이 _pending_widget_sync 우회 경로에
            등록되지 않은 채 _import_session_backup에서 직접 대입되고 있었음.
            STEP 2 영역에서 이미 instantiate된 text_area라 직접 수정 불가.
    · 수정: _BACKUP_KEYS 순회 중 위젯 key 두 개만 건너뛰고,
            pending dict에 등록 → 페이지 최상단 _pending_sync 처리에서 안전 적용.
    · 효과: 첫 클릭에 바로 복원 성공 (재클릭 불필요), 빨간 오류 박스 사라짐.

★ v2.0.7 미세 패치 (2026-05-23)
  [패치 L] 최종 모드 DOCX 표지 메타정보 제거
    · 원인: 최종 모드(비트 헤더 제거)로 출력해도 첫 페이지에
            "시나리오 / EPISODE N / 장르 / 기획·제작 / Series Engine vN · NN비트"
            5줄 메타가 그대로 박혀 나옴. 제작·연출·투자 전달용 각본에
            엔진 메타가 노출되는 문제.
    · 수정: make_series_docx_bytes 시그니처에 include_cover 인자 추가.
            기본값 True (기존 동작 유지). build_docx_download에서
            include_beat_headers=False(최종 모드)인 경우 자동으로
            include_cover=False 전달 → 표지 페이지 전체 스킵.
    · 효과: 최종 모드 출력 시 본문 1페이지에서 S#1부터 바로 시작.
            집필 모드(비트 헤더 포함)는 종전대로 표지 출력.

★ v2.0.8 미세 패치 (2026-05-23)
  [패치 M] 씬 번호 회별 리셋 — 시즌 통산 채번 차단
    · 원인: 「왕게임」 실측에서 EP2가 S#58, EP6이 S#215, EP8이 S#312로
            시즌 통산 채번이 발생. prompt.py 룰은 "에피소드 단위 연속"으로
            정상 명시되어 있으나, build_write_episode_beat_prompt 함수에
            구조적 결함이 있었음:
            - main.py L3079: 새 EP의 첫 비트(Beat 0) 집필 시 prev_beat_text로
              직전 EP 마지막 비트(EP-1, Beat 7)를 그대로 전달.
            - prompt.py L2655: prev_beat_text에서 마지막 씬 번호 추출 후
              "이 비트는 S#{N+1}부터 시작하라" 강제 블록을 생성.
            - 결과: EP1 Beat 7이 S#57에서 끝나면, EP2 Beat 0이 S#58부터
              시작하도록 모델에게 강제. EP가 진행될수록 통산 채번이 누적.
    · 수정:
            (1) build_write_episode_beat_prompt에 is_first_beat_of_episode
                인자 추가 (기본 False, 하위 호환).
            (2) is_first_beat_of_episode=True면 scene_continuation_block을
                "S#1부터 리셋" 강제 블록으로 대체.
            (3) prev_block 라벨도 "직전 EP{N-1} 마지막 부분 — 정서·상황
                연결용 참조 / 씬 번호는 새 EP에서 S#1부터 다시 시작"으로
                명시화. 모델이 분위기·잔향만 가져오고 씬 번호는 이어받지
                않도록 분리.
            (4) main.py 호출부에서 next_beat==0 조건으로 플래그 전달.
            (5) prompt.py [씬 번호 채번] 룰에 "EP 경계 리셋" 원칙 추가.
                잘못된 예 / 올바른 예 모두 EP 경계 케이스로 보강.
    · 효과: 새 EP의 첫 비트는 무조건 S#1부터 시작. 시즌 통산 채번 차단.
            「왕게임」 같은 누적형 누출 패턴 재발 방지.
"""

import streamlit as st
import anthropic
import io
import re
import json
from datetime import datetime

from profession_pack import build_multi_profession_block, detect_profession_category
from period_pack import build_period_block_auto

from prompt import (
    SYSTEM_PROMPT,
    GENRE_RULES,
    SEASON_BEATS_8,
    SEASON_BEATS_6,
    EPISODE_BEATS,
    BEAT_STRUCTURE_TYPES,
    build_locked_block,
    build_season_arc_prompt,
    build_character_expansion_prompt,
    build_episode_character_prompt,
    build_event_expansion_prompt,
    build_extract_elements_prompt,
    build_episode_plan_prompt,
    build_write_episode_beat_prompt,
    build_rewrite_prompt,
    build_structural_rewrite_prompt,
    summarize_episode_context,
    extract_from_creator_json_series,
    get_essence_check,
    sanitize_json_string,
    build_adaptive_context,
    JSON_OUTPUT_RULES,
)

# ═══════════════════════════════════════════════════════════
# ★ v1.9 패치 A — Writer Engine DOCX 빌더 자산 이식
# (분단 알고리즘 / PROP 메모 strip / INSERT 시스템)
# Writer Engine v3.x main.py의 헬퍼 블록을 그대로 가져와
# Series Engine 컨텍스트에서 동작하도록 통합.
# ═══════════════════════════════════════════════════════════

ENGINE_VERSION = "v2.0.8"
ENGINE_BUILD_DATE = "2026-05-23"


# ═══════════════════════════════════════════════════════════
# ★ v2.0 — 시즌 표현 누적 시스템
# 비트 집필 후 핵심 표현(20자 이상 연속 문장)을 자동 추출해
# 시즌 단위 누적 DB로 관리. 다음 비트 집필 시 3회 이상 누적된
# 표현은 모델에게 명시적 금지 지시.
# 정량 보고서 "AI Escape A17 위반 147건" 해결책.
# ═══════════════════════════════════════════════════════════

def extract_signature_expressions(text: str, min_len: int = 18, max_len: int = 80) -> list:
    """비트 집필 결과 텍스트에서 시그니처 표현을 추출.

    추출 대상:
    - 18자 이상 80자 이하의 문장 (한 문장 단위)
    - 시나리오 본문(지문/대사)만 — 메타·헤더·구분선 제외

    Returns:
        list[str] — 정규화된 표현 목록 (중복 제거)
    """
    import re as _re
    if not text:
        return []

    # 메타 영역 잘라내기 — '---' 또는 '═══' 이후는 메모
    body = text
    for sep in ['\n---\n', '\n═══', '\n━━━', '\n[BLOCK 2', '\n<WRITER_NOTES']:
        idx = body.find(sep)
        if idx > 0:
            body = body[:idx]

    # 줄 단위로 분리 후 메타·헤더 라인 제외
    lines = []
    for raw_line in body.split('\n'):
        line = raw_line.strip()
        if not line:
            continue
        # 씬 헤딩 제외
        if _re.match(r'^S#\d+\.', line):
            continue
        # 비트 헤더 / EP 헤더 제외
        if _re.match(r'^(Beat\s+\d+|EP\d+\s*[—\-]|EPISODE\s+\d+|═{3,}|─{3,}|={3,}|━{3,})', line):
            continue
        # 메타 prefix 제외 (간이판)
        if _re.match(r'^[\-•·*⭐★□]\s*(\*\*)?(서사동력|관객 심리|빌런|비밀|장르|본질|SCOPE|AI ESCAPE|fun_engine|absolute_goal|emotion_triggers|Plant|Payoff|핵심|모티프|맥거핀|보이스|액션 아이디어|비트)', line):
            continue
        # 캐릭터 + \t\t + 대사 형식 → 대사 부분만 추출
        m_dialogue = _re.match(r'^([가-힣]{1,5})\s*\t+(.+)$', line)
        if m_dialogue:
            line = m_dialogue.group(2).strip()
        lines.append(line)

    # 문장 단위 분리 (마침표·물음표·느낌표 뒤)
    text_joined = ' '.join(lines)
    sentences = _re.split(r'(?<=[.!?])\s+', text_joined)

    expressions = []
    seen = set()
    for sent in sentences:
        s = sent.strip()
        # 길이 필터
        if not (min_len <= len(s) <= max_len):
            continue
        # 정규화 — 끝의 구두점 제거 후 중복 검사
        norm = s.rstrip('.!?').strip()
        if norm in seen:
            continue
        seen.add(norm)
        expressions.append(s)

    return expressions


def update_season_expression_db(
    db: dict, new_text: str, ep_num: int, beat_num: int,
) -> dict:
    """시즌 표현 누적 DB 갱신.

    Args:
        db: 현재 누적 DB. { 표현 텍스트: 누적 횟수 }
        new_text: 방금 집필 완료된 비트 텍스트
        ep_num: EP 번호 (디버깅 용)
        beat_num: Beat 번호 (디버깅 용)

    Returns:
        갱신된 DB (입력 db 객체를 그대로 mutate 후 반환)
    """
    if db is None:
        db = {}
    expressions = extract_signature_expressions(new_text)
    for expr in expressions:
        norm = expr.strip().rstrip('.!?').strip()
        if not norm:
            continue
        db[norm] = db.get(norm, 0) + 1
    return db


def get_overused_expressions(db: dict, threshold: int = 3) -> list:
    """누적 N회 이상인 표현 목록 반환 (높은 순)."""
    if not db:
        return []
    items = [(e, c) for e, c in db.items() if c >= threshold]
    items.sort(key=lambda x: -x[1])
    return items


# ─────────────────────────────────────
# 지문(액션 라인) 자동 분단 헬퍼
# AI가 6~9문장을 한 문단에 몰아넣은 경우, 의미 비트 단위로 분단.
# AI의 시적 의도가 명확한 경우(짧은 단락)는 절대 건드리지 않음.
# ─────────────────────────────────────

# 분단 임계값
_ACTION_SPLIT_CHAR_THRESHOLD = 150
_ACTION_SPLIT_SENTENCE_THRESHOLD = 7
_ACTION_SPLIT_HARD_CHARS = 240
_ACTION_SPLIT_HARD_SENTENCES = 9


def _split_sentences(text: str):
    """한국어 지문을 문장 단위로 쪼갠다."""
    import re as _re
    parts = _re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in parts if s.strip()]


def _detect_paragraph_break_index(sentences: list) -> int:
    """문장 리스트에서 가장 자연스러운 분단 위치(인덱스)를 찾는다."""
    import re as _re
    n = len(sentences)
    if n < 4:
        return -1

    candidates = []
    for i in range(2, n - 1):
        cur = sentences[i]
        prev = sentences[i - 1]
        score = 0

        # 시간 압축 / 큰 상황 전환
        time_break_patterns = [
            r'수업이?\s*(끝나|진행)',
            r'시간이\s*(흐른|지난|경과)',
            r'(다음\s*날|이튿날|새벽|아침|저녁|밤)',
            r'몇\s*(분|시간|일)\s*(후|뒤)',
            r'(직후|잠시\s*후|곧)',
        ]
        for pat in time_break_patterns:
            if _re.search(pat, cur):
                score += 10
                break

        if _re.search(r'(혼자|홀로|단독)', cur):
            score += 6

        space_patterns = [
            r'^(긴\s|넓은\s|좁은\s|텅\s|빈\s|새|작은\s|커다란\s)',
            r'^(테이블|벽|창|문|바닥|천장|복도|골목)\s',
            r'^(아일랜드|카운터|책상|의자|침대|소파)\s',
            r'^[가-힣]+\s위에',
        ]
        prev_has_actor = bool(
            _re.search(r'[가-힣]{2,4}이\s', prev) or _re.search(r'[가-힣]{2,4}가\s', prev)
        )
        cur_is_space = any(_re.search(pat, cur) for pat in space_patterns)
        if prev_has_actor and cur_is_space:
            score += 5

        prev_actor = _re.match(r'^([가-힣]{2,4})(이|가)\s', prev)
        cur_actor = _re.match(r'^([가-힣]{2,4})(이|가)\s', cur)
        if prev_actor and cur_actor and prev_actor.group(1) != cur_actor.group(1):
            score += 4

        center_distance = abs(i - n // 2)
        position_bonus = max(0, 3 - center_distance)
        score += position_bonus

        if score >= 5:
            candidates.append((i, score))

    if not candidates:
        return -1
    candidates.sort(key=lambda x: -x[1])
    return candidates[0][0]


def _split_action_paragraph(text: str, _depth: int = 0) -> list:
    """지문 단락이 임계값을 넘으면 의미 비트 단위로 분할.

    ★ v1.9 안전 가드: 재귀 깊이 제한 + 동일 출력 감지로 무한 재귀 방지.
    """
    text = text.strip()
    if not text:
        return [text]

    # 재귀 깊이 제한 — 100단계 이상은 그대로 반환
    if _depth >= 100:
        return [text]

    char_len = len(text)
    sentences = _split_sentences(text)
    sent_count = len(sentences)

    triggered_by_length = char_len >= _ACTION_SPLIT_CHAR_THRESHOLD
    triggered_by_sentence = (char_len >= 100 and sent_count >= _ACTION_SPLIT_SENTENCE_THRESHOLD)

    if not (triggered_by_length or triggered_by_sentence):
        return [text]

    # 문장이 1개 이하면 분할 불가
    if sent_count <= 1:
        return [text]

    split_idx = _detect_paragraph_break_index(sentences)
    if split_idx < 0:
        if char_len < _ACTION_SPLIT_HARD_CHARS and sent_count < _ACTION_SPLIT_HARD_SENTENCES:
            return [text]
        split_idx = sent_count // 2

    # 분할 인덱스가 양 끝이면 분할 무효 — 그대로 반환
    if split_idx <= 0 or split_idx >= sent_count:
        return [text]

    part1 = ' '.join(sentences[:split_idx])
    part2 = ' '.join(sentences[split_idx:])

    # 분할 결과가 원본과 거의 동일하면 (분할 실패) 종료
    if len(part1) == 0 or len(part2) == 0:
        return [text]
    if part1 == text or part2 == text:
        return [text]

    return [part1] + _split_action_paragraph(part2, _depth + 1)


# ─────────────────────────────────────
# PROP CONTINUITY 메모 / 자가 검증 태그 자동 제거
# ─────────────────────────────────────

import re as _re_prop


def _strip_prop_state_memos(text: str) -> str:
    """텍스트에서 [소품 상태 / ...] 메모, GENRE_*_CHECK 태그 등을 제거."""
    if not text:
        return text

    pattern_codeblock = _re_prop.compile(
        r'```[^\n]*\n\[소품\s*상태[^\]]*\][\s\S]*?```',
        _re_prop.MULTILINE,
    )
    text = pattern_codeblock.sub('', text)

    pattern_inline = _re_prop.compile(
        r'\n*\[소품\s*상태[^\]]*\]\s*\n'
        r'(?:[\s]*[-•·][^\n]*\n?)+',
        _re_prop.MULTILINE,
    )
    text = pattern_inline.sub('\n', text)

    pattern_internal = _re_prop.compile(
        r'\n*\[?(?:INTERNAL|작가\s*노트|작가노트|소품\s*추적)[^\]]*\]?\s*\n'
        r'(?:\[소품\s*상태[^\]]*\]\s*\n)?'
        r'(?:[\s]*[-•·][^\n]*\n?)+',
        _re_prop.IGNORECASE | _re_prop.MULTILINE,
    )
    text = pattern_internal.sub('\n', text)

    for tag in ('GENRE_BOOSTER_CHECK', 'HELPER_CHARACTER_CHECK', 'GENRE_ESSENCE_CHECK'):
        pattern = _re_prop.compile(
            r'\n*<' + tag + r'>[\s\S]*?</' + tag + r'>\n*',
            _re_prop.IGNORECASE,
        )
        text = pattern.sub('\n', text)

    pattern_check_header = _re_prop.compile(
        r'\n*\[★?\s*비트\s*종료[^\]]*(?:GENRE_BOOSTER_CHECK|GENRE_ESSENCE_CHECK)[^\]]*\][\s\S]*?(?=\n\[|\nS#|\n$|\Z)',
        _re_prop.IGNORECASE,
    )
    text = pattern_check_header.sub('\n', text)

    text = _re_prop.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ─────────────────────────────────────
# INSERT 시스템 (화면 텍스트 표기)
# ─────────────────────────────────────

import re as _re_insert

_INSERT_LABEL_KEYWORDS = [
    '카톡', '메신저', '라인', '디스코드', '카카오톡',
    '문자', 'SMS', 'MMS',
    '이메일', '메일',
    '유튜브', 'YouTube', 'youtube', 'TV', '뉴스', '방송',
    'SNS', '인스타', '인스타그램', '페이스북', '트위터', 'X', '틱톡', 'DM',
    '검색', '구글', '네이버', '다음',
    '노트', '일기', '메모', '편지', '손글씨', '쪽지',
    '신문', '잡지', '기사', '헤드라인',
    '자막',
    '알림',
    '핸드폰', '핸드폰 화면', '폰', '폰 화면', '화면',
]


def _is_insert_label(text: str) -> bool:
    text = text.strip()
    if not (text.startswith('[') and ']' in text):
        return False
    label_match = _re_insert.match(r'^\[([^\]]+)\]', text)
    if not label_match:
        return False
    label_inner = label_match.group(1)
    return any(kw in label_inner for kw in _INSERT_LABEL_KEYWORDS)


def _parse_insert_blocks(text: str) -> list:
    """여러 줄 텍스트를 받아 INSERT 블록과 일반 텍스트로 분리."""
    if not text or not text.strip():
        return []

    lines = text.split('\n')
    items = []
    i = 0
    n = len(lines)
    accumulated_action = []

    def flush_action():
        if accumulated_action:
            joined = '\n'.join(accumulated_action).strip()
            if joined:
                items.append({'type': 'action', 'data': joined})
            accumulated_action.clear()

    while i < n:
        line = lines[i]
        line_stripped = line.strip()

        # 형식 A: INSERT — / INSERT - / INSERT:
        if _re_insert.match(r'^INSERT\s*[—\-:]', line_stripped, _re_insert.IGNORECASE):
            flush_action()
            header = line_stripped
            body_lines = []
            i += 1
            while i < n:
                bl = lines[i].strip()
                if _re_insert.match(r'^\[/INSERT\]?$', bl, _re_insert.IGNORECASE):
                    i += 1
                    break
                if not bl:
                    j = i + 1
                    while j < n and not lines[j].strip():
                        j += 1
                    if j >= n:
                        i = j
                        break
                    next_line = lines[j].strip()
                    if _re_insert.match(r'^\[/INSERT\]?$', next_line, _re_insert.IGNORECASE):
                        i = j + 1
                        break
                    if not _re_insert.match(r"^['\"\u2018\u2019\u201C\u201D]", next_line):
                        i = j
                        break
                    i += 1
                    continue
                body_lines.append(bl)
                i += 1
            items.append({
                'type': 'insert_block',
                'data': {'header': header, 'body': body_lines},
            })
            continue

        if _is_insert_label(line_stripped):
            flush_action()
            items.append({'type': 'insert_label', 'data': line_stripped})
            i += 1
            continue

        if _re_insert.match(r'^\[/INSERT\]?$', line_stripped, _re_insert.IGNORECASE):
            i += 1
            continue

        accumulated_action.append(line)
        i += 1

    flush_action()
    return items


def _parse_insert_label(text: str) -> tuple:
    """형식 B 라벨 한 줄을 (label, body)로 분리."""
    m = _re_insert.match(r'^(\[[^\]]+\])\s*(.*)$', text.strip())
    if m:
        return m.group(1), m.group(2).strip()
    return text, ""


# ──────────────────────────────────────────────
# Page Config
# ──────────────────────────────────────────────
st.set_page_config(
    page_title="BLUE JEANS · Series Engine",
    page_icon="👖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ──────────────────────────────────────────────
# ★ v2.0 — Sidebar Engine Info (Writer/Creator Engine과 동일 톤)
# ──────────────────────────────────────────────
with st.sidebar:
    st.markdown(f"""
    <div style="padding:12px;background:#F0F2FF;border-radius:8px;border-left:3px solid #191970;font-family:'Pretendard',sans-serif;">
        <div style="font-size:.72rem;color:#191970;font-weight:700;letter-spacing:.05em;margin-bottom:4px;">ENGINE INFO</div>
        <div style="font-size:1.05rem;font-weight:700;color:#191970;">Series Engine</div>
        <div style="font-size:1.25rem;font-weight:900;color:#FFCB05;background:#191970;padding:2px 8px;border-radius:4px;display:inline-block;margin-top:4px;">
            {ENGINE_VERSION}
        </div>
        <div style="font-size:.7rem;color:#666;margin-top:8px;">
            Build: {ENGINE_BUILD_DATE}<br>
            Creator Engine v2.5.3+ 호환<br>
            Writer Engine 양식 정합
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.caption("버전이 최신인지 확인하세요.")

    # ★ v2.0 — 시즌 표현 누적 현황 (집필 중일 때만)
    _expr_db = st.session_state.get("season_expression_db", {}) or {}
    if _expr_db:
        _overused = [(e, c) for e, c in _expr_db.items() if c >= 3]
        _total = len(_expr_db)
        _over_count = len(_overused)
        st.markdown(f"""
        <div style="padding:10px;background:#FFF9E6;border-radius:8px;border-left:3px solid #FFCB05;margin-top:12px;font-family:'Pretendard',sans-serif;">
            <div style="font-size:.7rem;color:#7A6500;font-weight:700;letter-spacing:.05em;margin-bottom:4px;">★ 시즌 표현 누적 (v2.0)</div>
            <div style="font-size:.78rem;color:#1A1A2E;line-height:1.4;">
                추적 표현: <b>{_total}</b>개<br>
                3회 이상 누적: <b style="color:#D32F2F">{_over_count}</b>개
            </div>
            <div style="font-size:.62rem;color:#888;margin-top:4px;">다음 비트 집필 시 자동 차단</div>
        </div>
        """, unsafe_allow_html=True)

# ──────────────────────────────────────────────
# Custom CSS — Creator Engine 디자인 시스템 통일
# ──────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #191970;
    --y: #FFCB05;
    --bg: #F7F7F5;
    --card: #FFFFFF;
    --card-border: #E2E2E0;
    --t: #1A1A2E;
    --r: #D32F2F;
    --g: #2EC484;
    --dim: #8E8E99;
    --light-bg: #EEEEF6;
    --serif: 'Paperlogy', 'Noto Serif KR', 'Georgia', serif;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

/* ── 기본 타이포 ── */
html, body, [class*="css"] {
    font-family: var(--body);
    color: var(--t);
    -webkit-font-smoothing: antialiased;
}

/* ══ 라이트모드 강제 ══ */
.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important;
    color: var(--t) !important;
}
.stMarkdown, .stText, .stCode {
    color: var(--t) !important;
}
h1, h2, h3, h4, h5, h6 { color: var(--navy) !important; font-family: var(--heading) !important; }
p, span, label, div, li { color: inherit; }

/* ── 사이드바 숨김 ── */
section[data-testid="stSidebar"] { display: none; }

/* ══ 입력 위젯 ══ */
.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-size: 0.9rem !important;
    padding: 0.6rem 0.8rem !important;
    transition: border-color 0.2s;
}
.stTextInput input:focus, .stTextArea textarea:focus,
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(25,25,112,0.08) !important;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder,
[data-testid="stTextInput"] input::placeholder, [data-testid="stTextArea"] textarea::placeholder {
    color: var(--dim) !important;
    font-size: 0.85rem !important;
}
/* selectbox */
.stSelectbox > div > div, [data-baseweb="select"] > div, [data-baseweb="select"] input {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border-color: var(--card-border) !important;
    border-radius: 8px !important;
}
[data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"], [role="option"] {
    background-color: var(--card) !important;
    color: var(--t) !important;
}
[role="option"]:hover { background-color: var(--light-bg) !important; }
/* label */
.stTextInput label, .stTextArea label, .stSelectbox label, .stRadio label {
    color: var(--t) !important;
    font-weight: 600 !important;
    font-size: 0.82rem !important;
    margin-bottom: 0.3rem !important;
}

/* ══ 버튼 ══ */
.stButton > button {
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    background-color: var(--card) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    padding: 0.5rem 1.2rem !important;
    transition: all 0.2s;
}
.stButton > button:hover {
    border-color: var(--navy) !important;
    box-shadow: 0 2px 8px rgba(25,25,112,0.08) !important;
}
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background-color: var(--y) !important;
    color: var(--navy) !important;
    border-color: var(--y) !important;
    font-weight: 700 !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background-color: #E8B800 !important;
    box-shadow: 0 2px 12px rgba(255,203,5,0.3) !important;
}

/* ══ Expander ══ */
.stExpander, details, details summary {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1px solid var(--card-border) !important;
    border-radius: 8px !important;
}
details[open] > div { background-color: var(--card) !important; }
.stExpander summary, .stExpander summary span { color: var(--t) !important; }

/* ══ Alert ══ */
.stAlert { color: var(--t) !important; border-radius: 8px !important; }

/* ══ 내부 컨테이너 투명 ══ */
[data-testid="stVerticalBlock"], [data-testid="stHorizontalBlock"],
[data-testid="stColumn"] { background-color: transparent !important; }

/* ══ Metric ══ */
[data-testid="stMetric"] { background-color: var(--card) !important; color: var(--t) !important; }
[data-testid="stMetric"] label { color: var(--dim) !important; }
.stCaption, small { color: var(--dim) !important; }

/* ═══════════════════════════════════
   브랜딩 & 커스텀 컴포넌트
   ═══════════════════════════════════ */

.header {
    font-size: 0.85rem;
    font-weight: 700;
    color: var(--navy);
    letter-spacing: 0.15em;
    margin-bottom: 0;
    font-family: var(--heading);
}

.brand-title {
    font-size: 2.6rem;
    font-weight: 900;
    color: var(--navy);
    font-family: var(--display);
    letter-spacing: -0.02em;
    margin-bottom: 0.15rem;
    position: relative;
    display: inline-block;
}
.brand-title::after {
    content: '';
    position: absolute;
    bottom: 2px;
    left: 0;
    width: 100%;
    height: 4px;
    background: var(--y);
    border-radius: 2px;
}

.sub {
    font-size: 0.7rem;
    color: var(--dim);
    letter-spacing: 0.15em;
    margin-top: 0.5rem;
    margin-bottom: 1.5rem;
}

/* ── 카드 ── */
.card {
    background: var(--card);
    border: 1px solid var(--card-border);
    border-radius: 10px;
    padding: 1.2rem;
    margin-bottom: 0.8rem;
    box-shadow: 0 1px 4px rgba(0,0,0,0.03);
    transition: all 0.2s;
}
.card:hover {
    border-color: var(--navy);
    box-shadow: 0 3px 12px rgba(25,25,112,0.07);
    transform: translateY(-1px);
}

/* ── 콜아웃 ── */
.callout {
    background: var(--light-bg);
    border-left: 4px solid var(--navy);
    padding: 0.9rem 1.1rem;
    margin: 0.5rem 0;
    border-radius: 0 8px 8px 0;
    font-size: 0.88rem;
    color: var(--t);
}

/* ── 섹션 라벨 ── */
.cl {
    color: var(--navy);
    font-weight: 700;
    font-size: 0.72rem;
    letter-spacing: 0.03em;
    margin-bottom: 0.3rem;
    text-transform: uppercase;
}

/* ── 정보 블록 ── */
.ri {
    background: var(--light-bg);
    border-radius: 8px;
    padding: 0.9rem 1rem;
    margin-bottom: 0.5rem;
    font-size: 0.88rem;
    line-height: 1.6;
}
.rl {
    color: var(--navy);
    font-weight: 700;
    font-size: 0.72rem;
    letter-spacing: 0.02em;
}

/* ── 뱃지 ── */
.badge {
    display: inline-block;
    padding: 0.2rem 0.6rem;
    border-radius: 5px;
    font-size: 0.7rem;
    font-weight: 600;
}
.b-done { background: var(--g); color: #fff; }
.b-run  { background: var(--y); color: var(--navy); }
.b-not  { background: #E8E8F0; color: var(--dim); }

/* ── 노란 섹션 헤더 ── */
.section-header {
    background: var(--y);
    color: var(--navy);
    padding: 0.6rem 1rem;
    border-radius: 6px;
    font-weight: 800;
    font-size: 1rem;
    font-family: var(--heading);
    margin: 1.5rem 0 0.8rem 0;
    display: flex;
    justify-content: space-between;
    align-items: center;
}
.section-header .en {
    font-family: var(--display);
    font-size: 0.75rem;
    font-weight: 700;
    letter-spacing: 0.05em;
    opacity: 0.7;
}

/* ══ Stepper ══ */
.stepper {
    display: flex;
    align-items: center;
    justify-content: center;
    margin: 1.5rem 0 2rem 0;
    gap: 0;
}
.step-wrap {
    display: flex;
    flex-direction: column;
    align-items: center;
}
.step-circle {
    width: 34px;
    height: 34px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.78rem;
    font-weight: 700;
    flex-shrink: 0;
    border: 2px solid transparent;
    transition: all 0.2s;
}
.step-circle.active {
    background: var(--y);
    color: var(--navy);
    border-color: var(--y);
    box-shadow: 0 0 0 4px rgba(255,203,5,0.2);
}
.step-circle.done {
    background: var(--g);
    color: #fff;
    border-color: var(--g);
}
.step-circle.upcoming {
    background: #EDEDF0;
    color: var(--dim);
    border-color: #D8D8E0;
}
.step-label {
    font-size: 0.6rem;
    margin-top: 0.35rem;
    text-align: center;
    width: 55px;
    font-weight: 500;
}
.step-label.active { color: var(--navy); font-weight: 700; }
.step-label.done { color: var(--g); font-weight: 600; }
.step-label.upcoming { color: var(--dim); }
.step-line {
    width: 30px;
    height: 2px;
    margin: 0 2px;
    flex-shrink: 0;
    border-radius: 1px;
}
.step-line.done { background: var(--g); }
.step-line.upcoming { background: #D8D8E0; }

/* ══ 에피소드 탭 ══ */
.ep-tabs {
    display: flex;
    flex-wrap: wrap;
    gap: 0.4rem;
    margin: 0.8rem 0;
}
.ep-tab {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    padding: 0.35rem 0.9rem;
    border-radius: 6px;
    font-size: 0.78rem;
    font-weight: 700;
    font-family: var(--body);
    transition: all 0.2s;
}
.ep-tab.done { background: var(--g); color: #fff; }
.ep-tab.current { background: var(--y); color: var(--navy); box-shadow: 0 0 0 3px rgba(255,203,5,0.2); }
.ep-tab.pending { background: #EDEDF0; color: var(--dim); }

/* ══ 비트 프로그레스 ══ */
.beat-row {
    display: flex;
    gap: 0.35rem;
    flex-wrap: wrap;
    margin: 0.8rem 0;
}
.beat-dot {
    width: 2.2rem;
    height: 2.2rem;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.7rem;
    font-weight: 700;
    font-family: var(--body);
    transition: all 0.2s;
    border: 2px solid transparent;
}
.beat-dot.done { background: var(--g); color: #fff; border-color: var(--g); }
.beat-dot.current { background: var(--y); color: var(--navy); border-color: var(--y); box-shadow: 0 0 0 3px rgba(255,203,5,0.2); }
.beat-dot.pending { background: #EDEDF0; color: var(--dim); border-color: #D8D8E0; }

/* ══ 프로그레스 바 ══ */
.progress-track {
    flex: 1;
    background: #E8E8F0;
    border-radius: 4px;
    height: 8px;
    margin: 0 0.5rem;
}
.progress-fill {
    height: 100%;
    background: var(--y);
    border-radius: 4px;
    transition: width 0.3s;
}
</style>
""", unsafe_allow_html=True)


# ──────────────────────────────────────────────
# 모델 & 토큰 설정
# ──────────────────────────────────────────────
MODEL_WRITE = "claude-opus-4-6"       # 집필 (비트 쓰기, 다시 쓰기) — 최고 품질
MODEL_PLAN  = "claude-sonnet-4-6"    # 구조 작업 (시즌 아크, 씬 플랜, 요소 추출) — 비용 효율
MAX_TOKENS_ARC = 8000
MAX_TOKENS_PLAN = 8000
MAX_TOKENS_BEAT = 16000
MAX_TOKENS_REWRITE = 16000


# ──────────────────────────────────────────────
# 브랜드 헤더 — Creator Engine과 동일 패턴
# ──────────────────────────────────────────────
st.markdown(
    '<div style="text-align:center;padding:1rem 0 0 0">'
    '<div class="header">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>'
    '<div class="brand-title">SERIES ENGINE</div>'
    '<div class="sub">Y O U N G &nbsp; · &nbsp; V I N T A G E &nbsp; · &nbsp; F R E E &nbsp; · &nbsp; I N N O V A T I V E</div>'
    '</div>',
    unsafe_allow_html=True
)
st.caption(f"집필: {MODEL_WRITE} · 구조: {MODEL_PLAN}")


# ──────────────────────────────────────────────
# API 클라이언트
# ──────────────────────────────────────────────
@st.cache_resource
def get_client():
    api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("🔑 `ANTHROPIC_API_KEY`를 Streamlit Secrets에 설정하세요.")
        st.stop()
    return anthropic.Anthropic(api_key=api_key)


def stream_response(system: str, user_prompt: str, max_tokens: int, model: str = ""):
    """Claude API 스트리밍 호출. model 미지정 시 WRITE 모델 사용."""
    use_model = model or MODEL_WRITE
    client = get_client()
    collected = []
    placeholder = st.empty()
    with client.messages.stream(
        model=use_model,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_prompt}],
    ) as stream:
        for text in stream.text_stream:
            collected.append(text)
            placeholder.markdown("".join(collected))
    return "".join(collected)


# ──────────────────────────────────────────────
# 세션 스테이트
# ──────────────────────────────────────────────
DEFAULTS = {
    "inputs": {
        "title": "",  # ★ v2.0.5 — 작품 제목 (백업 파일명·사이드바·PDF 푸터)
        "logline": "", "intention": "", "gns": "",
        "characters": "", "world": "", "structure": "",
        "scenes": "", "treatment": "", "tone": "",
    },
    "num_episodes": 8,
    "duration": 50,
    "genre": "범죄/스릴러",
    "season_arc": "",
    "story_elements": "",
    "expanded_characters": "",
    "expanded_events": "",
    "episode_plans": {},
    "episode_beats": {},
    "beat_structure_types": {},
    "episode_summaries": {},
    "ep_characters": {},
    "locked_items": [],
    "open_items": [],
    "producer_notes_write": "",
    "season_expression_db": {},  # v2.0
    "monitoring_feedback": "",  # v2.0.4 — 외부 모니터링 의견
    "showrunner_notes": "",     # v2.0.4 — 쇼러너 / 작가 노트
}

for k, v in DEFAULTS.items():
    if k not in st.session_state:
        if isinstance(v, dict):
            st.session_state[k] = dict(v)
        else:
            st.session_state[k] = v


# ★ v1.8 패치: pending 위젯 동기화 처리
# 백업 복원이나 Creator JSON 로드 후 rerun된 시점에 위젯 key를 안전하게 갱신.
# 위젯이 그려지기 전(이 시점)이라 직접 수정 가능.
_pending_sync = st.session_state.pop("_pending_widget_sync", None)
if _pending_sync:
    for _wk, _wv in _pending_sync.items():
        st.session_state[_wk] = _wv


# ──────────────────────────────────────────────
# Stepper (5단계)
# ──────────────────────────────────────────────
STEPS = [
    ("input",   "자료입력"),
    ("arc",     "시즌아크"),
    ("plan",    "씬플랜"),
    ("write",   "집필"),
    ("export",  "다운로드"),
]


def get_current_step() -> int:
    if len(st.session_state["episode_beats"]) > 0:
        return 3
    if len(st.session_state["episode_plans"]) > 0:
        return 2
    if st.session_state["season_arc"]:
        return 1
    return 0


def get_done_step() -> int:
    ne = st.session_state["num_episodes"]
    total_beats = ne * 8
    done_beats = len(st.session_state["episode_beats"])
    if done_beats >= total_beats:
        return 4
    if done_beats > 0:
        return 3
    if len(st.session_state["episode_plans"]) > 0:
        return 2
    if st.session_state["season_arc"]:
        return 1
    return -1


def render_stepper():
    current_idx = get_current_step()
    done_idx = get_done_step()
    html_parts = []
    for i, (key, label) in enumerate(STEPS):
        if i <= done_idx and i < current_idx:
            state = "done"
        elif i == current_idx:
            state = "active"
        else:
            state = "upcoming"
        circle = "✓" if state == "done" else str(i + 1)
        html_parts.append(
            f'<div class="step-wrap">'
            f'<div class="step-circle {state}">{circle}</div>'
            f'<div class="step-label {state}">{label}</div>'
            f'</div>'
        )
        if i < len(STEPS) - 1:
            line_state = "done" if i < current_idx and i <= done_idx else "upcoming"
            html_parts.append(f'<div class="step-line {line_state}"></div>')
    st.markdown('<div class="stepper">' + ''.join(html_parts) + '</div>', unsafe_allow_html=True)


render_stepper()


# ──────────────────────────────────────────────
# 유틸리티
# ──────────────────────────────────────────────

def bk(ep: int, beat: int) -> str:
    return f"{ep}_{beat}"


def _get_locked_block() -> str:
    """현재 세션의 LOCKED/OPEN 항목으로 프롬프트 주입 블록 생성."""
    locked = st.session_state.get("locked_items", [])
    open_items = st.session_state.get("open_items", [])
    if not locked and not open_items:
        return ""
    return build_locked_block(locked, open_items)


def get_all_episode_text(ep: int) -> str:
    parts = []
    for b in range(8):
        key = bk(ep, b)
        if key in st.session_state["episode_beats"]:
            parts.append(st.session_state["episode_beats"][key])
    return "\n\n".join(parts)


def get_full_season_text() -> str:
    ne = st.session_state["num_episodes"]
    parts = []
    for ep in range(1, ne + 1):
        ep_text = get_all_episode_text(ep)
        if ep_text:
            parts.append(f"{'='*60}\nEPISODE {ep}\n{'='*60}\n\n{ep_text}")
    return "\n\n\n".join(parts)


def build_txt_download(text: str, filename: str):
    st.download_button(
        label=f"📄 {filename}",
        data=text.encode("utf-8"),
        file_name=filename,
        mime="text/plain",
    )


def build_docx_download(text: str, filename: str, title: str = "",
                        include_beat_headers: bool = False):
    """v1.9 — Writer Engine 양식 빌더 호출 래퍼.

    기존 평문 빌더 자리를 한국 시나리오 표준 양식 빌더로 교체.
    text가 시즌 전체(`============================================================
EPISODE N` 마커 포함)인지 단일 EP 텍스트인지 자동 판단해
    어댑터 함수로 위임. 메타데이터(장르/부수/비트 수)는 session_state에서 자동 추출.
    
    ★ v2.0.3 — include_beat_headers 기본 False (최종 모드).
    ★ v2.0.7 — 최종 모드(include_beat_headers=False)인 경우 표지 페이지
               자동 제거 (include_cover=False 전달). 집필 모드는 표지 유지.
    """
    try:
        # session_state에서 메타데이터 추출 (있으면)
        _genre = st.session_state.get("genre", "") or ""
        _num_eps = st.session_state.get("num_episodes", 0) or 0
        _beats_count = len(st.session_state.get("episode_beats", {}) or {})

        # ★ v2.0.7 — 최종 모드면 표지 페이지 제거
        _include_cover = bool(include_beat_headers)

        # 시즌 전체 마커 감지
        if "\nEPISODE " in text and ("=" * 30) in text:
            data = make_series_docx_bytes(
                text, title=title, mode="season",
                genre=_genre, num_episodes=_num_eps,
                beats_done_count=_beats_count,
                include_beat_headers=include_beat_headers,
                include_cover=_include_cover,
            )
        else:
            data = make_series_docx_bytes(
                text, title=title, mode="episode",
                genre=_genre, num_episodes=_num_eps,
                beats_done_count=_beats_count,
                include_beat_headers=include_beat_headers,
                include_cover=_include_cover,
            )

        st.download_button(
            label=f"📘 {filename}",
            data=data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ImportError:
        st.warning("python-docx가 설치되지 않았습니다. TXT 다운로드를 이용하세요.")
    except Exception as e:
        st.error(f"DOCX 생성 중 오류가 발생했습니다: {e}")


# ═══════════════════════════════════════════════════════════
# ★ v1.9 패치 A — 한국 시나리오 표준 양식 DOCX 빌더
# Writer Engine v3.x의 make_docx_bytes 자산을 Series Engine
# 컨텍스트로 이식. 양식 스타일 7종 + 메타 차단 + INSERT 시스템.
# ═══════════════════════════════════════════════════════════

# Series Engine 전용 메타 차단 키워드 (Writer Engine 100+ 키워드에 추가)
# v1.8 「수비니어샵」 SEASON1_FULL 분석으로 식별된 본문 누출 패턴
_SERIES_META_PREFIX_PATTERNS = [
    # 비트 단위 메타 헤더
    "비트 요약", "비트요약",
    "SCOPE 검증", "SCOPE검증",
    "5단계 추적", "5단계추적",
    "작동한 장르 장치", "작동한 장르 장치 목록",
    "핵심 요소 추적", "핵심요소추적",
    "관객 심리", "관객심리",
    "비밀 경제", "비밀경제",
    "빌런 추적", "빌런추적", "Villain 4Q", "Villain 추적",
    "LOCKED 검증", "LOCKED검증",
    "서사동력 추적", "서사동력추적", "서사동력",
    "보이스 점검", "보이스점검",
    "장르 드라이브 5점", "장르드라이브 5점",
    "액션 아이디어 전진", "액션 아이디어",
    "AI ESCAPE 점검", "AI ESCAPE", "AI Escape", "AI 탈출",
    "장르 강제 체크", "장르강제 체크", "Genre Enforcement",
    "본질 검증", "본질검증",
    "장르 드리프트 체크", "장르 드리프트",
    "fun_engine 작동", "fun_engine",
    "B-Story 진행", "B-Story",
    "BJND 4축", "BJND",
    # 본질 3중 선언 (Creator Engine v2.5.5+ 호환)
    "absolute_goal", "emotion_triggers", "fun_engine",
    # 장르 강제 체크 4축 (Series Engine 「수비니어샵」 64회 등장 — 매 비트마다)
    "정보 비대칭", "정보비대칭",
    "시계 장치", "시계장치",
    "도덕선 이동", "도덕선이동",
    "물리적 위협", "물리적위협",
    # 씬 플랜 단위 메타
    "맥거핀 1", "맥거핀 2", "맥거핀 3", "맥거핀 4",
    "맥거핀1", "맥거핀2", "맥거핀3", "맥거핀4",
    "맥거핀",
    "캐릭터 비밀", "캐릭터비밀",
    "핵심 장소", "핵심장소",
    "모티프",
    "열린 질문", "열린질문",
    "Dramatic Irony",
    "Zeigarnik",
    # 시리즈 단위 비밀
    "시즌 비밀", "중간 비밀", "에피소드 비밀",
    # Writer Engine 공통 메타 (중복이지만 안전망)
    "writer_notes", "plant_payoff_tag", "plant_payoff",
    "scene_meta", "quality_check",
    "Plant:", "Plant/Payoff", "Payoff:", "Payoff :",
    "Plant 유지", "Plant유지", "Payoff 회수", "Payoff회수",
    "서브플롯", "서브플롯 진행",
    "캐릭터 전술", "캐릭터전술", "캐릭터 아크",
    "비밀",
    "장르 드라이브", "장르드라이브",
    # 장르 장치 snake_case (Writer Engine 90개 — Series에서도 누출)
    "premise_engine", "comic_contradiction", "character_comic_flaw",
    "comic_escalation", "line_surprise", "status_comedy",
    "timing_precision", "callback_payoff", "scene_comic_engine",
    "joke_density",
    "fear_anticipation", "uncertainty", "sensory_unease",
    "threat_design", "dread_pacing", "violation_of_safety",
    "image_residue", "vulnerability", "false_relief",
    "terror_escalation",
    "information_asymmetry", "escalation", "clock_device",
    "suspense_peak", "plot_twist", "investigator_obstacle",
    "villain_intelligence", "moral_ambiguity", "red_herring",
    "irreversible_stakes",
    "action_spark", "physical_choreography", "setpiece_scale",
    "hero_signature", "obstacle_escalation", "stakes_personal",
    "counter_attack", "low_point", "final_confrontation",
    "kinetic_rhythm",
    "longing_distance", "touch_hesitation", "romantic_specificity",
    "emotional_subtext", "miscommunication", "emotional_reversal",
    "vulnerability_moment", "physical_chemistry", "obstacle_internal",
    "payoff_emotional",
    "world_rule", "tech_showcase", "awe_moment", "info_drip",
    "human_anchor", "rule_consequence", "visual_wonder",
    "scale_shift", "philosophical_stakes", "discovery_rhythm",
    "magic_rule", "mythic_echo", "threshold_crossing",
    "wonder_image", "sacrifice_price", "prophecy_twist",
    # 추가 본문 누출 후보 (Series Engine 진단으로 식별)
    "pressure_escalation", "threat_visibility_control",
    "desire_origin",
    # 한자 / 특수 메타
    "수미상관", "수미 상관",
    "Cold Opening", "콜드 오프닝",
]


def make_series_docx_bytes(
    text: str,
    title: str = "",
    mode: str = "season",
    genre: str = "",
    num_episodes: int = 0,
    beats_done_count: int = 0,
    fact_based: bool = False,
    historical: bool = False,
    historical_type: str = "",
    include_beat_headers: bool = False,
    include_cover: bool = True,
) -> bytes:
    """시리즈 시나리오 DOCX — 한국 표준 시나리오 서식.

    Series Engine episode_beats 텍스트를 받아 Writer Engine 양식으로 렌더링.

    Parameters
    ----------
    text : str
        Series Engine 본문 텍스트 (시즌 전체 또는 단일 EP).
    title : str
        커버 페이지에 표시할 제목 (작품명 또는 EP명).
    mode : "season" | "episode"
        season: 시즌 표지 + EP 분리 페이지 + 비트 헤더 처리.
        episode: 단일 EP 표지 + 비트 헤더 처리.
    genre : str
        장르 표시 (커버 페이지 부제).
    num_episodes : int
        시즌 총 EP 수 (커버 페이지 부제).
    beats_done_count : int
        완성 비트 수 (커버 페이지 진행도).
    fact_based / historical / historical_type :
        면책 자막 페이지 자동 삽입 조건.
    include_beat_headers : bool, default False
        ★ v2.0.3 — 비트헤더 출력 여부.
        True  → 집필 모드 (Beat 0~7 헤더 표시, 작가 검토용)
        False → 최종 모드 (비트 헤더 제거, 제작·연출 전달용) — 기본값
    include_cover : bool, default True
        ★ v2.0.7 — 표지 페이지 출력 여부.
        True  → 표지 페이지 출력 (시나리오 / 제목 / 부제 / 기획·제작 / 엔진버전)
        False → 표지 페이지 생략 (S#1부터 바로 시작 — 최종 전달용)
    """
    import re
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Mm, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from io import BytesIO

    doc = DocxDocument()

    # ── 페이지 설정 (A4, 20mm 마진) ──
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # ── 기본 스타일: 함초롬바탕 10pt ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "함초롬바탕"
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(2)
    style_normal.paragraph_format.space_before = Pt(0)
    rpr = style_normal.element.rPr
    if rpr is None:
        rpr = style_normal.element.makeelement(qn('w:rPr'), {})
        style_normal.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), '함초롬바탕')

    def _set_eastasia_font(rpr_elem, font_name='함초롬바탕'):
        rf = rpr_elem.find(qn('w:rFonts'))
        if rf is None:
            rf = rpr_elem.makeelement(qn('w:rFonts'), {})
            rpr_elem.append(rf)
        rf.set(qn('w:eastAsia'), font_name)

    # ── 양식 스타일 7종 + Series 전용 2종 ──

    # [1] 씬번호
    style_scene = doc.styles.add_style('씬번호', WD_STYLE_TYPE.PARAGRAPH)
    style_scene.base_style = doc.styles['Normal']
    style_scene.font.name = '함초롬바탕'
    style_scene.font.size = Pt(11)
    style_scene.font.bold = True
    style_scene.paragraph_format.space_before = Pt(24)
    style_scene.paragraph_format.space_after = Pt(6)
    style_scene.paragraph_format.line_spacing = 1.5
    _set_eastasia_font(style_scene.element.get_or_add_rPr())

    # [2] 대사
    style_dialogue = doc.styles.add_style('대사', WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue.base_style = doc.styles['Normal']
    style_dialogue.font.name = '함초롬바탕'
    style_dialogue.font.size = Pt(10)
    style_dialogue.font.bold = True
    style_dialogue.paragraph_format.left_indent = Cm(1.25)
    style_dialogue.paragraph_format.space_before = Pt(8)
    style_dialogue.paragraph_format.space_after = Pt(2)
    style_dialogue.paragraph_format.line_spacing = 1.5
    _set_eastasia_font(style_dialogue.element.get_or_add_rPr())

    # [3] 대사연속
    style_dialogue_cont = doc.styles.add_style('대사연속', WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue_cont.base_style = style_dialogue
    style_dialogue_cont.paragraph_format.space_before = Pt(0)
    style_dialogue_cont.paragraph_format.space_after = Pt(0)

    # [4] 지문
    style_action = doc.styles.add_style('지문', WD_STYLE_TYPE.PARAGRAPH)
    style_action.base_style = doc.styles['Normal']
    style_action.font.name = '함초롬바탕'
    style_action.font.size = Pt(10)
    style_action.font.bold = False
    style_action.paragraph_format.space_before = Pt(2)
    style_action.paragraph_format.space_after = Pt(2)
    _set_eastasia_font(style_action.element.get_or_add_rPr())

    # [5] INSERT 헤더
    style_insert_header = doc.styles.add_style('인서트헤더', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_header.base_style = doc.styles['Normal']
    style_insert_header.font.name = '함초롬바탕'
    style_insert_header.font.size = Pt(9)
    style_insert_header.font.bold = True
    style_insert_header.paragraph_format.left_indent = Cm(2.55)
    style_insert_header.paragraph_format.space_before = Pt(8)
    style_insert_header.paragraph_format.space_after = Pt(2)
    _set_eastasia_font(style_insert_header.element.get_or_add_rPr())

    # [6] INSERT 본문
    style_insert_body = doc.styles.add_style('인서트본문', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_body.base_style = doc.styles['Normal']
    style_insert_body.font.name = '함초롬바탕'
    style_insert_body.font.size = Pt(10)
    style_insert_body.font.italic = True
    style_insert_body.paragraph_format.left_indent = Cm(2.55)
    style_insert_body.paragraph_format.space_before = Pt(2)
    style_insert_body.paragraph_format.space_after = Pt(2)
    style_insert_body.paragraph_format.line_spacing = 1.4
    _set_eastasia_font(style_insert_body.element.get_or_add_rPr())

    # [7] INSERT 라벨
    style_insert_label = doc.styles.add_style('인서트라벨', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_label.base_style = doc.styles['Normal']
    style_insert_label.font.name = '함초롬바탕'
    style_insert_label.font.size = Pt(10)
    style_insert_label.paragraph_format.left_indent = Cm(1.42)
    style_insert_label.paragraph_format.space_before = Pt(4)
    style_insert_label.paragraph_format.space_after = Pt(4)
    _set_eastasia_font(style_insert_label.element.get_or_add_rPr())

    # [Series-1] EP 분리 헤더 (Series 전용)
    style_ep_heading = doc.styles.add_style('에피소드헤더', WD_STYLE_TYPE.PARAGRAPH)
    style_ep_heading.base_style = doc.styles['Normal']
    style_ep_heading.font.name = '함초롬바탕'
    style_ep_heading.font.size = Pt(20)
    style_ep_heading.font.bold = True
    style_ep_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_ep_heading.paragraph_format.space_before = Pt(48)
    style_ep_heading.paragraph_format.space_after = Pt(24)
    _set_eastasia_font(style_ep_heading.element.get_or_add_rPr())

    # [Series-2] 비트 구분 헤더 (Series 전용)
    style_beat_heading = doc.styles.add_style('비트헤더', WD_STYLE_TYPE.PARAGRAPH)
    style_beat_heading.base_style = doc.styles['Normal']
    style_beat_heading.font.name = '함초롬바탕'
    style_beat_heading.font.size = Pt(10)
    style_beat_heading.font.bold = True
    style_beat_heading.font.italic = True
    style_beat_heading.paragraph_format.space_before = Pt(18)
    style_beat_heading.paragraph_format.space_after = Pt(8)
    _set_eastasia_font(style_beat_heading.element.get_or_add_rPr())

    # ── 헬퍼 함수 ──
    def add_text(s, bold=False, size=None, color=None, align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(s)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        if bold:
            r.bold = True
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color
        return p

    def add_scene_heading(s):
        p = doc.add_paragraph(style='씬번호')
        r = p.add_run(s)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_dialogue(char_name, parenthetical, line, continuation=False):
        if continuation:
            p = doc.add_paragraph(style='대사연속')
            speaker_part = "\t\t"
        else:
            p = doc.add_paragraph(style='대사')
            speaker_part = f"{char_name}\t\t"

        body_parts = []
        if parenthetical:
            body_parts.append((f"({parenthetical}) ", True))
        if line:
            import re as _re
            chunks = _re.split(r'(\([^()]*\))', line)
            for chunk in chunks:
                if not chunk:
                    continue
                if chunk.startswith('(') and chunk.endswith(')'):
                    body_parts.append((chunk, True))
                else:
                    body_parts.append((chunk, False))

        r_speaker = p.add_run(speaker_part)
        r_speaker.font.name = "함초롬바탕"
        _set_eastasia_font(r_speaker._element.get_or_add_rPr())

        for txt, is_paren in body_parts:
            r = p.add_run(txt)
            r.font.name = "함초롬바탕"
            _set_eastasia_font(r._element.get_or_add_rPr())
            if is_paren:
                r.bold = False
        return p

    def add_blank_line():
        p = doc.add_paragraph(style='지문')
        r = p.add_run("")
        r.font.name = "함초롬바탕"
        r.font.size = Pt(10)
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_insert_block(header, body_lines):
        doc.add_paragraph("")
        first_p = doc.add_paragraph(style='인서트헤더')
        r = first_p.add_run(header.strip())
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        for line in body_lines:
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='인서트본문')
            r = p.add_run(line)
            r.font.name = "함초롬바탕"
            r.italic = True
            _set_eastasia_font(r._element.get_or_add_rPr())
        close_p = doc.add_paragraph(style='인서트헤더')
        cr = close_p.add_run('[/INSERT]')
        cr.font.name = "함초롬바탕"
        _set_eastasia_font(cr._element.get_or_add_rPr())
        doc.add_paragraph("")
        return first_p

    def add_insert_label(s):
        label, body = _parse_insert_label(s)
        p = doc.add_paragraph(style='인서트라벨')
        r_label = p.add_run(label + ' ')
        r_label.font.name = "함초롬바탕"
        r_label.font.size = Pt(9)
        r_label.bold = True
        _set_eastasia_font(r_label._element.get_or_add_rPr())
        if body:
            r_body = p.add_run(body)
            r_body.font.name = "함초롬바탕"
            r_body.italic = True
            _set_eastasia_font(r_body._element.get_or_add_rPr())
        return p

    def add_action(s):
        items = _parse_insert_blocks(s)
        first_p = None
        for item in items:
            if item['type'] == 'insert_block':
                p = add_insert_block(item['data']['header'], item['data']['body'])
            elif item['type'] == 'insert_label':
                p = add_insert_label(item['data'])
            else:
                sub_paragraphs = _split_action_paragraph(item['data'])
                p = None
                for sub in sub_paragraphs:
                    sp = doc.add_paragraph(style='지문')
                    r = sp.add_run(sub)
                    r.font.name = "함초롬바탕"
                    _set_eastasia_font(r._element.get_or_add_rPr())
                    if p is None:
                        p = sp
            if first_p is None:
                first_p = p
        return first_p

    def add_episode_heading(ep_num):
        """EP 분리 페이지 헤더."""
        doc.add_page_break()
        p = doc.add_paragraph(style='에피소드헤더')
        r = p.add_run(f"EPISODE {ep_num}")
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_beat_heading(s):
        """비트 구분 헤더 (예: 'EP1 — Beat 0. Cold Opening').
        
        ★ v2.0.3 — include_beat_headers=False (기본)인 경우 출력 안 함.
        최종 시나리오에는 작업 메타정보가 들어가면 안 됨.
        """
        if not include_beat_headers:
            return None  # 최종 모드 — 비트헤더 출력 안 함
        p = doc.add_paragraph(style='비트헤더')
        r = p.add_run(s)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    # ── 커버 페이지 ──
    # ★ v2.0.7 — include_cover=False 시 표지 전체 스킵 (제작·연출 전달용)
    if include_cover:
        for _ in range(6):
            doc.add_paragraph("")
        add_text("시나리오", size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph("")
        add_text(title or "<무제>", bold=True, size=Pt(24), align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph("")

        # 부제 (시즌 + 부수 + 장르)
        subtitle_parts = []
        if mode == "season" and num_episodes:
            subtitle_parts.append(f"시즌 1 — {num_episodes}부작")
        if genre:
            subtitle_parts.append(genre)
        if subtitle_parts:
            add_text(" · ".join(subtitle_parts), size=Pt(12),
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     color=RGBColor(0x66, 0x66, 0x66))
            doc.add_paragraph("")

        doc.add_paragraph("")
        add_text("기획/제작 | 블루진픽처스", size=Pt(10),
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        progress = f"  ·  {beats_done_count}비트" if beats_done_count else ""
        add_text(f"Series Engine {ENGINE_VERSION}{progress}",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        doc.add_page_break()

    # ── 면책 자막 페이지 ──
    _need_disclaimer = fact_based or (
        historical and (
            "팩션" in (historical_type or "") or "퓨전" in (historical_type or "")
            or "faction" in (historical_type or "").lower()
            or "fusion" in (historical_type or "").lower()
        )
    )
    if _need_disclaimer:
        for _ in range(10):
            doc.add_paragraph("")
        add_text("본 작품에 등장하는 인물, 단체, 지명, 상호, 사건은",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text("모두 허구이며, 실존하는 것과 관련이 있더라도",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text("극적 구성을 위해 각색되었습니다.",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph("")
        add_text("All characters, organizations, places, and events in this work",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        add_text("are fictional. Any resemblance to actual persons or events is",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        add_text("dramatized for narrative purposes.",
                 size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x8E, 0x8E, 0x99))
        doc.add_page_break()

    # ── 본문 파싱 ──
    # PROP 메모 / GENRE_*_CHECK 태그 strip
    text = _strip_prop_state_memos(text)

    # 정규식 패턴
    heading_re = re.compile(
        r'^S?#?\d*\.?\s*(INT\.|EXT\.|INT\./EXT\.)\s*(.+)',
        re.IGNORECASE,
    )
    char_re = re.compile(
        r'^\s{2,}([가-힣a-zA-Z\s]{1,15}?)\s*'
        r'(?:\((V\.O\.|O\.S\.|CONT\'D|cont\'d|v\.o\.|o\.s\.)\))?\s*$',
        re.IGNORECASE,
    )
    inline_dialogue_re = re.compile(
        r'^([가-힣a-zA-Z\s]{1,15}?)\s*'
        r'(?:\((V\.O\.|O\.S\.|CONT\'D|cont\'d|v\.o\.|o\.s\.)\))?\s*'
        r'\t{1,}\s*(?:\(([^)]*)\)\s*)?(.+)',
        re.IGNORECASE,
    )
    paren_re = re.compile(r'^\s{2,}\((.+?)\)\s*$')
    divider_re = re.compile(r'^(═{3,}|─{3,}|={3,}|---)')

    # EP 분리 마커 (Series 전용): "EPISODE N" 라인 + 위/아래 ===== 구분선
    ep_marker_re = re.compile(r'^EPISODE\s+(\d+)\s*$')
    # 비트 헤더 마커: "EP1 — Beat 0. Cold Opening" 또는 "Beat 1 — ..."
    beat_marker_re = re.compile(r'^(?:EP\d+\s*[—\-]\s*)?Beat\s+\d+', re.IGNORECASE)

    # ── 메타 차단 필터 ──
    # ★ v1.9.1 — 정규식 prefix에 마크다운 bold(**) 마커 흡수 추가
    # AI가 출력 시 ` - ⭐ **본질 검증**: ` 같은 마크다운 강조를 섞어 쓰는 경우가
    # 「수비니어샵」 실측에서 44건 누출됐기 때문에 ** 마커 양쪽 흡수가 필수.
    # ★ v1.9.2 — catastrophic backtracking 방지
    # nested quantifier (`(?:...+\**)*`) 가 `═` 39개 같은 입력에서 폭발했음.
    # prefix 부분을 단일 character class 한 번에 흡수하는 형태로 평탄화.
    META_PREFIX_PATTERNS = _SERIES_META_PREFIX_PATTERNS
    META_LINE_RE = re.compile(
        r'^[\s•·\-*⭐★─═□☑✓①②③④⑤⑥⑦⑧⑨⑩]{0,40}\s*\**\s*'
        r'(' + '|'.join(re.escape(p) for p in META_PREFIX_PATTERNS) + r')'
        r'\**(?:\s|[:\-—.(]|$)',
        re.IGNORECASE,
    )
    META_DEV_NOTATION_RE = re.compile(
        r'\((?:Beat\s*\d+|S#\s*\d+|전체|전반|후반)[^)]*'
        r'(?:plant|payoff|→|->)[^)]*\)',
        re.IGNORECASE,
    )
    META_DEV_COMMENT_RE = re.compile(
        r'^[\s•·\-*]+.*?\(S#\s*\d+(?:[/,]\s*\d+)*\)\s*(?:—|-|–)\s*'
        r'(?:미공개|미등장|미해결|공개|폭로|열린|유지|부재)',
        re.IGNORECASE,
    )
    META_IRONY_COMMENT_RE = re.compile(
        r'\(관객\s*[OX].*?(?:유진|주인공|캐릭터|클레어|관객)\s*[OX][^)]*\)\s*(?:—|-|–)\s*'
        r'(?:미공개|미등장|미해결|공개|유지)',
    )
    META_CHARACTER_SUMMARY_RE = re.compile(
        r'^[\s•·\-*]+\s*([가-힣]+(?:[·∙・]\s*[가-힣]+)*(?:\s*커플)?)\s*:',
    )
    META_BULLET_DEV_RE = re.compile(
        r'^[\s•·\-*]+.*?\((?:Beat\s*\d+|S#\s*\d+)[^)]*\)\s*:',
        re.IGNORECASE,
    )
    # Series 전용: "□ XXX: YES/NO" 형태 (장르 강제 체크 4축)
    META_CHECKBOX_RE = re.compile(
        r'^[\s]*□\s*[가-힣A-Za-z]',
    )
    # Series 전용: snake_case 그룹 라벨 ("absolute_goal:" 등)
    META_SNAKECASE_RE = re.compile(
        r'^[\s•·\-*]*\s*([a-z]+(?:_[a-z]+)+)\s*[:\-—]',
    )
    # Series 전용: "맥거핀 N:" 또는 "캐릭터의 비밀:" 식 콜론 메타
    # ★ v1.9.1 — 마크다운 ** 마커 흡수
    META_COLON_HEADER_RE = re.compile(
        r'^[\s•·\-*]*\**\s*(맥거핀\s*\d*|캐릭터\s*비밀|핵심\s*장소|모티프|열린\s*질문|'
        r'(?:[가-힣]{2,4})\s*의?\s*비밀|시즌\s*비밀|중간\s*비밀|에피소드\s*비밀)\**\s*:',
    )

    def is_meta_line(s):
        if not s:
            return False
        if META_LINE_RE.match(s):
            return True
        if re.match(r'^[\s•·\-*⭐★─═]+\s*([a-z]+_[a-z_]+)', s):
            return True
        if META_BULLET_DEV_RE.match(s):
            return True
        if META_CHARACTER_SUMMARY_RE.match(s):
            return True
        if META_DEV_NOTATION_RE.search(s):
            return True
        if META_DEV_COMMENT_RE.match(s):
            return True
        if META_IRONY_COMMENT_RE.search(s):
            return True
        if META_CHECKBOX_RE.match(s):
            return True
        if META_SNAKECASE_RE.match(s):
            return True
        if META_COLON_HEADER_RE.match(s):
            return True
        return False

    # ── 캐릭터명 형식 붕괴 복구 (v2.0.1 강화) ──
    # AI가 다양한 형식으로 출력한 캐릭터+대사 분리를 한 단락으로 복구.
    # 처리 케이스:
    #   A) "캐릭터\n\n대사"     — 빈 줄 사이
    #   B) "캐릭터\n대사"        — 빈 줄 없이 연속  
    #   C) "캐릭터\n(괄호)\n대사" — 행동 지시 분리
    #   D) "캐릭터 (V.O.)\n대사" — V.O./O.S./F 같은 변형
    # 모두 "캐릭터\t\t대사" 또는 "캐릭터\t\t(괄호) 대사" 단일 라인으로 통합.
    # Series 전용 — 캐릭터명 자동 추출
    _series_char_pattern = re.compile(
        r'^([가-힣]{2,8}|[A-Z][a-zA-Z]+|[가-힣]{2,8}\s*#?\d*|남자|여자|위반자|손님|국과수[가-힣\s]*|팀장|반장)'
        r'\s*(\(\s*[VOF]\.?[OS]?\.?\s*\)|\(\s*F\s*\))?$'
    )

    def _is_likely_char_name(s):
        """줄 단독으로 등장한 텍스트가 캐릭터명일 가능성 판정."""
        s = s.strip()
        if not s or len(s) > 20:
            return False
        # 끝의 V.O. / O.S. / F 등 제거 후 본체 검사
        s_clean = re.sub(r'\s*\(\s*[VOF]\.?[OS]?\.?\s*\)\s*$', '', s).strip()
        s_clean = re.sub(r'\s*\(\s*F\s*\)\s*$', '', s_clean).strip()
        if not s_clean or len(s_clean) > 12:
            return False
        # 한글 2~8자 또는 영문 (대문자 시작)
        if not re.match(r'^([가-힣]{2,8}|[A-Z][a-zA-Z]+)(\s*#?\d+)?$', s_clean):
            # 특수 캐릭터명 화이트리스트
            if s_clean not in {'위반자', '손님', '남자', '여자', '팀장', '반장',
                              '운영자', '국과수 김 주임', '국과수 담당자', '집행자'}:
                return False
        # 흔한 비-캐릭터 명사 배제
        excluded = {'그리고', '하지만', '그러나', '그래서', '그러면', '왜냐하면',
                    '결국', '갑자기', '마침내', '여전히', '아직도', '이제는',
                    '그날', '그때', '그곳', '여기', '저기', '거기',
                    '인사동', '서울', '한국', '오늘', '내일', '어제',
                    '카운터', '카메라', '서랍', '계단', '복도', '문턱',
                    '아이템', '매뉴얼', '계약서', '정산서', '영수증', '청구서', '처방전'}
        if s_clean in excluded:
            return False
        return True

    _broken_lines = text.split("\n")
    _fixed_lines = []
    _j = 0
    while _j < len(_broken_lines):
        _cur = _broken_lines[_j].strip()
        
        # 캐릭터명 후보가 아니면 그대로 통과
        if not _is_likely_char_name(_cur):
            _fixed_lines.append(_broken_lines[_j])
            _j += 1
            continue
        
        # 다음 비어있지 않은 줄 찾기
        _next_idx = _j + 1
        while _next_idx < len(_broken_lines) and not _broken_lines[_next_idx].strip():
            _next_idx += 1
        
        # 다음 콘텐츠 줄이 없으면 그대로
        if _next_idx >= len(_broken_lines):
            _fixed_lines.append(_broken_lines[_j])
            _j += 1
            continue
        
        _next_content = _broken_lines[_next_idx].strip()
        
        # 다음 줄이 씬번호/비트헤더/EP헤더면 캐릭터명만 있는 거 — 합치지 않음
        if (_next_content.startswith("S#") or _next_content.startswith("EP") or
            _next_content.startswith("Beat") or _next_content.startswith("===") or
            _next_content.startswith("━━") or _next_content.startswith("───")):
            _fixed_lines.append(_broken_lines[_j])
            _j += 1
            continue
        
        # 다음 줄도 캐릭터명이면 합치지 않음
        if _is_likely_char_name(_next_content):
            _fixed_lines.append(_broken_lines[_j])
            _j += 1
            continue
        
        # 케이스 C: 다음이 (괄호) 행동 지시 + 그 다음이 실제 대사
        if (_next_content.startswith("(") and _next_content.endswith(")")
                and len(_next_content) <= 40):
            # 그 다음 비어있지 않은 줄 찾기
            _after_paren = _next_idx + 1
            while _after_paren < len(_broken_lines) and not _broken_lines[_after_paren].strip():
                _after_paren += 1
            if (_after_paren < len(_broken_lines) and
                _broken_lines[_after_paren].strip() and
                not _is_likely_char_name(_broken_lines[_after_paren].strip())):
                _fixed_lines.append(
                    f"{_cur}\t\t{_next_content} {_broken_lines[_after_paren].strip()}"
                )
                _j = _after_paren + 1
                continue
        
        # 일반 케이스: 캐릭터 + 다음 줄 대사 합치기
        _fixed_lines.append(f"{_cur}\t\t{_next_content}")
        _j = _next_idx + 1
    lines = _fixed_lines

    # ── 본문 변환 루프 ──
    i = 0
    prev_block_type = None  # "scene" | "action" | "dialogue" | "insert" | "ep" | "beat"

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # ── Series 전용 EP 분리 마커 ──
        # 위 라인이 ===== 구분선, 현 라인이 "EPISODE N", 다음 라인이 ===== 면 EP 분리
        if ep_marker_re.match(stripped):
            ep_match = ep_marker_re.match(stripped)
            ep_num = ep_match.group(1)
            add_episode_heading(ep_num)
            prev_block_type = "ep"
            i += 1
            continue

        # ── Series 전용 비트 헤더 마커 ──
        # "EP1 — Beat 0. Cold Opening" 또는 "═══════ Beat 1 ═══════" 식
        if beat_marker_re.match(stripped):
            add_beat_heading(stripped)
            prev_block_type = "beat"
            i += 1
            continue

        # ── WRITER_NOTES 마커 블록 스킵 ──
        if "<WRITER_NOTES_BEGIN>" in stripped or "WRITER_NOTES_BEGIN" in stripped:
            i += 1
            while i < len(lines):
                if ("<WRITER_NOTES_END>" in lines[i] or "WRITER_NOTES_END" in lines[i]):
                    i += 1
                    break
                i += 1
            continue

        if "<SPACE_DIVERSITY_CHECK>" in stripped or "SPACE_DIVERSITY_CHECK" in stripped:
            i += 1
            while i < len(lines):
                if "</SPACE_DIVERSITY_CHECK>" in lines[i] or (
                    "SPACE_DIVERSITY_CHECK" in lines[i] and "</" in lines[i]
                ):
                    i += 1
                    break
                i += 1
            continue

        if stripped.startswith("[BLOCK 2:") or stripped.startswith("[BLOCK 2 "):
            i = len(lines)
            continue
        if stripped == "[BLOCK 1: 시나리오 본문]" or stripped.startswith("[BLOCK 1:"):
            i += 1
            continue
        if stripped.startswith("━━━"):
            i += 1
            continue

        # 구분선 스킵 (--- / ═══ / === 등)
        # ★ v1.9.1 — 보수적 스킵 변경
        # 기존: divider 만나면 무조건 다음 헤딩까지 모두 스킵 → '--- TITLE ---' 같은
        #       정상 본문 다음 줄들도 사라지는 부작용
        # 개선: 다음 줄이 명백한 메모 신호(내부 메모/비트 요약/snake_case 등)일 때만
        #       블록 스킵. 아니면 구분선만 한 줄 버리고 다음 줄부터 정상 처리.
        if divider_re.match(stripped):
            # 다음 비-빈 줄 확인
            _peek_j = i + 1
            while _peek_j < len(lines) and not lines[_peek_j].strip():
                _peek_j += 1
            _next_line = lines[_peek_j].strip() if _peek_j < len(lines) else ""

            # 다음 줄이 명백한 메모 신호인지 판정
            _is_memo_block = False
            if _next_line:
                if "내부 메모" in _next_line:
                    _is_memo_block = True
                elif "비트 요약" in _next_line or "비트요약" in _next_line:
                    _is_memo_block = True
                elif is_meta_line(_next_line):
                    _is_memo_block = True
                elif re.match(r'^[\s•·\-*]*\**[a-z]+_[a-z]+', _next_line):
                    _is_memo_block = True

            if _is_memo_block:
                # 메모 블록 — 다음 헤딩까지 스킵
                i += 1
                while i < len(lines):
                    memo_line = lines[i].strip()
                    if heading_re.match(memo_line):
                        break
                    if ep_marker_re.match(memo_line):
                        break
                    if beat_marker_re.match(memo_line):
                        break
                    if "Beat " in memo_line and "—" in memo_line:
                        break
                    i += 1
                continue
            else:
                # 단순 구분선 — 그 줄만 버림
                i += 1
                continue

        # "내부 메모" 블록 스킵
        if "내부 메모" in stripped:
            i += 1
            while i < len(lines):
                memo_line = lines[i].strip()
                if heading_re.match(memo_line):
                    break
                if ep_marker_re.match(memo_line):
                    break
                if beat_marker_re.match(memo_line):
                    break
                i += 1
            continue

        # 개별 메타 라인 차단
        if is_meta_line(stripped):
            i += 1
            continue

        # "═" 또는 "Beat N — ..." 형태의 Beat 헤더 (메타 표기)
        # — beat_marker_re는 이미 위에서 처리되었으므로 여기는 보조 스킵
        if stripped.startswith("═"):
            i += 1
            continue

        # ── 씬 헤딩 ──
        m = heading_re.match(stripped)
        if m:
            add_scene_heading(stripped)
            prev_block_type = "scene"
            i += 1
            continue

        # ── 인라인 대사 ──
        im = inline_dialogue_re.match(stripped)
        if im:
            char_name = im.group(1).strip()
            vo_marker = im.group(2) or ""
            inline_paren = im.group(3) or ""
            inline_text = im.group(4).strip()
            if vo_marker:
                char_name = f"{char_name} ({vo_marker})"
            if prev_block_type in ("action", "insert"):
                add_blank_line()
            add_dialogue(char_name, inline_paren, inline_text)
            prev_block_type = "dialogue"
            i += 1
            continue

        # ── 캐릭터명 + 들여쓰기 대사 ──
        cm = char_re.match(line)
        if cm:
            char_name = cm.group(1).strip()
            vo_marker = cm.group(2) or ""
            if vo_marker:
                char_name = f"{char_name} ({vo_marker})"
            parenthetical = ""
            dialogue_lines = []
            if prev_block_type in ("action", "insert"):
                add_blank_line()
            i += 1

            if i < len(lines):
                pm = paren_re.match(lines[i])
                if pm:
                    parenthetical = pm.group(1)
                    i += 1

            while i < len(lines):
                dl = lines[i]
                ds = dl.strip()
                if not ds:
                    break
                if heading_re.match(ds):
                    break
                if ep_marker_re.match(ds):
                    break
                if beat_marker_re.match(ds):
                    break
                if char_re.match(dl):
                    break
                if inline_dialogue_re.match(ds):
                    break
                dialogue_lines.append(ds)
                i += 1

            if dialogue_lines:
                merged_parts = []
                current_dialogue = []
                for dl in dialogue_lines:
                    dl_s = dl.strip()
                    if (dl_s.startswith("(") and dl_s.endswith(")")
                            and len(dl_s) > 2):
                        if current_dialogue:
                            merged_parts.append(("dialogue", " ".join(current_dialogue)))
                            current_dialogue = []
                        merged_parts.append(("action", dl_s))
                    else:
                        current_dialogue.append(dl_s)
                if current_dialogue:
                    merged_parts.append(("dialogue", " ".join(current_dialogue)))

                first = True
                for part_type, part_text in merged_parts:
                    if part_type == "dialogue":
                        if first:
                            add_dialogue(char_name, parenthetical, part_text)
                            parenthetical = ""
                            first = False
                        else:
                            add_dialogue(char_name, "", part_text, continuation=True)
                    else:
                        add_blank_line()
                        add_action(part_text)
                        add_blank_line()
                        first = True
                if first:
                    add_dialogue(char_name, parenthetical, "")
            else:
                add_dialogue(char_name, parenthetical, "")
            prev_block_type = "dialogue"
            continue

        # ── 일반 지문 ──
        if prev_block_type == "dialogue":
            add_blank_line()
        add_action(stripped)
        prev_block_type = "action"
        i += 1

    # ── 푸터 ──
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"BLUE JEANS SERIES ENGINE {ENGINE_VERSION} · BLUE JEANS PICTURES")
    fr.font.size = Pt(7)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()
    add_text(f"© 2026 BLUE JEANS PICTURES · Series Engine {ENGINE_VERSION}",
             size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x8E, 0x8E, 0x99))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════

# ──────────────────────────────────────────────
# 모드 선택 (집필 / 리라이트)
# ──────────────────────────────────────────────
mode = st.radio(
    "모드",
    ["✍️ 집필", "🔄 리라이트"],
    horizontal=True,
    help="집필: 시즌 아크 → 씬 플랜 → 비트 집필 | 리라이트: 기존 비트를 프로듀서 노트·캐스팅 반영하여 구조적 재작성",
)

# ══════════════════════════════════════════════
# 🔄 리라이트 모드 (독립형 — 세션 비의존)
# ══════════════════════════════════════════════
if mode == "🔄 리라이트":
    st.markdown(
        '<div class="section-header">🔄 구조적 리라이트 <span class="en">STRUCTURAL REWRITE · NETFLIX MODE</span></div>',
        unsafe_allow_html=True,
    )
    st.caption(
        "기존 대본 + 기획서 + 프로듀서/캐스팅/모니터링 노트를 붙여넣어 구조적으로 재작성합니다. "
        "집필 모드와 독립적으로 작동합니다."
    )

    # ══════════════════════════════════════════════
    # ★ v2.0.5 — 리라이트 자동 로더 (Creator JSON + Series JSON)
    # ══════════════════════════════════════════════
    def _rw_build_subsequent_summary(beats_dict, from_ep):
        """후속 EP 비트들을 헤더 + 첫 300자 스니펫으로 압축. 토큰 폭주 방지."""
        if not beats_dict:
            return ""
        parts = []
        try:
            sorted_keys = sorted(beats_dict.keys(), key=lambda k: tuple(int(x) for x in k.split("_")))
        except (ValueError, IndexError):
            sorted_keys = list(beats_dict.keys())
        for k in sorted_keys:
            try:
                ep_part, bt_part = k.split("_")
                ep_int = int(ep_part)
                if ep_int <= from_ep:
                    continue
                v = beats_dict.get(k, "")
                if not v:
                    continue
                beat_name = ""
                if int(bt_part) < len(EPISODE_BEATS):
                    beat_name = EPISODE_BEATS[int(bt_part)].get("name", "")
                snippet = v.strip()[:300].replace("\n", " ")
                parts.append(f"[EP{ep_int} Beat{bt_part} — {beat_name}] {snippet}...")
            except (ValueError, IndexError):
                continue
        return "\n\n".join(parts)

    def _rw_sync_beat_from_loaded():
        """EP/Beat 셀렉터 변경 시 호출. 로드된 Series JSON에서 해당 비트 텍스트를 꺼내 pending에 적재."""
        loaded_beats = st.session_state.get("_rw_loaded_beats")
        if not loaded_beats:
            return
        ep = st.session_state.get("rw_ep_num", 1)
        bt = st.session_state.get("rw_beat_num", 0)
        key_str = f"{ep}_{bt}"
        beat_text = loaded_beats.get(key_str, "")
        subsequent_text = _rw_build_subsequent_summary(loaded_beats, ep)
        loaded_plans = st.session_state.get("_rw_loaded_plans", {})
        plan_text = loaded_plans.get(str(ep), "")

        pending = st.session_state.get("_pending_widget_sync", {}) or {}
        pending["rw_beat_text"] = beat_text
        pending["rw_ep_plan"] = plan_text
        pending["rw_subsequent"] = subsequent_text
        st.session_state["_pending_widget_sync"] = pending

    with st.expander("⚡ JSON 자동 로드 (Creator Engine / Series Engine)", expanded=False):
        st.markdown(
            '<div class="small-meta">'
            '<b>Creator Engine JSON</b>: 기획서 5칸(로그라인·GNS·캐릭터·세계관·톤) + LOCKED 자동 채움.<br>'
            '<b>Series Engine JSON</b>: 시즌 아크 · 씬 플랜 · 비트 본문 · 후속 EP를 EP/Beat 선택에 따라 자동 갱신.<br>'
            '외부 대본 리라이트는 이 단계 건너뛰고 아래 칸을 직접 채우세요.'
            '</div>',
            unsafe_allow_html=True,
        )

        col_rwj1, col_rwj2 = st.columns(2)

        # ── Creator Engine JSON ──
        with col_rwj1:
            st.markdown("**📋 Creator Engine JSON**")
            rw_creator_file = st.file_uploader(
                "Creator Engine JSON 파일",
                type=["json"],
                key="rw_creator_uploader",
                label_visibility="collapsed",
            )
            rw_load_creator_btn = st.button(
                "📂 Creator JSON 적용",
                use_container_width=True,
                disabled=(rw_creator_file is None),
                key="rw_load_creator_btn",
            )
            if rw_creator_file is not None:
                st.caption(f"선택됨: `{rw_creator_file.name}`")

            if rw_load_creator_btn and rw_creator_file is not None:
                try:
                    raw = rw_creator_file.read().decode("utf-8")
                    creator_data = json.loads(raw)
                    loaded = extract_from_creator_json_series(creator_data)

                    pending = st.session_state.get("_pending_widget_sync", {}) or {}

                    if loaded.get("logline"):
                        pending["rw_logline"] = loaded["logline"]
                    if loaded.get("gns"):
                        pending["rw_gns"] = loaded["gns"]
                    if loaded.get("characters"):
                        pending["rw_characters"] = loaded["characters"]
                    if loaded.get("world"):
                        pending["rw_world"] = loaded["world"]
                    if loaded.get("tone"):
                        pending["rw_tone"] = loaded["tone"]

                    # 작품 제목도 inputs.title에 주입 (사이드바·파일명용)
                    loaded_title = (loaded.get("title") or "").strip()
                    if loaded_title:
                        st.session_state["inputs"]["title"] = loaded_title
                        pending["input_title"] = loaded_title

                    locked_ext = loaded.get("locked_5_extended", "")
                    if locked_ext:
                        pending["rw_locked"] = locked_ext

                    if pending:
                        st.session_state["_pending_widget_sync"] = pending

                    meta = creator_data.get("_meta", {})
                    ce_ver = meta.get("engine_version", "?")
                    title = loaded.get("title", "(무제)")
                    filled = sum(1 for f in ["logline", "gns", "characters", "world", "tone"] if loaded.get(f))

                    st.success(
                        f"✅ Creator Engine {ce_ver} 로드 완료.\n\n"
                        f"**프로젝트**: {title}\n\n"
                        f"**기획서 필드**: {filled}개 / 5칸 채움"
                    )
                    st.rerun()
                except json.JSONDecodeError as e:
                    st.error(f"JSON 파싱 실패: {e}")
                except Exception as e:
                    st.error(f"로드 중 오류: {e}")

        # ── Series Engine JSON ──
        with col_rwj2:
            st.markdown("**📺 Series Engine JSON**")
            rw_series_file = st.file_uploader(
                "Series Engine 백업 JSON",
                type=["json"],
                key="rw_series_uploader",
                label_visibility="collapsed",
            )
            rw_load_series_btn = st.button(
                "📂 Series JSON 적용",
                use_container_width=True,
                disabled=(rw_series_file is None),
                key="rw_load_series_btn",
            )
            if rw_series_file is not None:
                st.caption(f"선택됨: `{rw_series_file.name}`")

            if rw_load_series_btn and rw_series_file is not None:
                try:
                    raw = rw_series_file.read().decode("utf-8")
                    series_data = json.loads(raw)
                    sess = series_data.get("session", {})
                    meta = series_data.get("_meta", {})

                    if not sess:
                        st.error("session 키가 없습니다. Series Engine 백업이 맞는지 확인하세요.")
                    else:
                        pending = st.session_state.get("_pending_widget_sync", {}) or {}

                        loaded_genre = sess.get("genre", "")
                        if loaded_genre and loaded_genre in GENRE_RULES:
                            pending["rw_genre"] = loaded_genre

                        # 작품 제목 복원 (inputs.title 또는 _meta.project_title)
                        sess_inputs = sess.get("inputs", {}) or {}
                        loaded_title = (sess_inputs.get("title") or "").strip() or (meta.get("project_title") or "").strip()
                        if loaded_title:
                            st.session_state["inputs"]["title"] = loaded_title
                            pending["input_title"] = loaded_title

                        exp_chars = sess.get("expanded_characters", "")
                        if exp_chars:
                            pending["rw_characters"] = exp_chars

                        season_arc = sess.get("season_arc", "")
                        if season_arc:
                            pending["rw_season_arc"] = season_arc

                        locked_list = sess.get("locked_items", []) or []
                        open_list = sess.get("open_items", []) or []
                        if locked_list:
                            pending["rw_locked"] = "\n".join(locked_list)
                        if open_list:
                            pending["rw_open"] = "\n".join(open_list)

                        prod_notes = sess.get("producer_notes_write", "")
                        if prod_notes:
                            pending["rw_producer"] = prod_notes

                        beats_dict = sess.get("episode_beats", {}) or {}
                        plans_dict = sess.get("episode_plans", {}) or {}
                        st.session_state["_rw_loaded_beats"] = beats_dict
                        st.session_state["_rw_loaded_plans"] = plans_dict

                        cur_ep = st.session_state.get("rw_ep_num", 1)
                        cur_bt = st.session_state.get("rw_beat_num", 0)
                        key_str = f"{cur_ep}_{cur_bt}"
                        if key_str in beats_dict:
                            pending["rw_beat_text"] = beats_dict[key_str]
                        if str(cur_ep) in plans_dict:
                            pending["rw_ep_plan"] = plans_dict[str(cur_ep)]

                        subseq_text = _rw_build_subsequent_summary(beats_dict, cur_ep)
                        if subseq_text:
                            pending["rw_subsequent"] = subseq_text

                        if pending:
                            st.session_state["_pending_widget_sync"] = pending

                        se_ver = meta.get("engine_version", "?")
                        saved_at = meta.get("saved_at", "?")
                        eps_planned = meta.get("episodes_planned", 0)
                        beats_written = meta.get("beats_written", 0)
                        title_display = loaded_title or "(제목 없음)"

                        st.success(
                            f"✅ Series Engine {se_ver} 백업 로드 완료.\n\n"
                            f"**작품**: {title_display}\n\n"
                            f"**저장 시각**: {saved_at}\n\n"
                            f"**장르**: {loaded_genre} / **씬 플랜**: {eps_planned}개 EP / **집필 비트**: {beats_written}개\n\n"
                            f"EP/Beat 셀렉터를 바꾸면 해당 비트 본문이 자동으로 따라옵니다."
                        )
                        st.rerun()
                except json.JSONDecodeError as e:
                    st.error(f"JSON 파싱 실패: {e}")
                except Exception as e:
                    st.error(f"로드 중 오류: {e}")

        if st.session_state.get("_rw_loaded_beats"):
            n_beats = len(st.session_state["_rw_loaded_beats"])
            n_plans = len(st.session_state.get("_rw_loaded_plans", {}))
            st.info(f"📺 Series JSON 로드됨: 비트 {n_beats}개, 씬 플랜 {n_plans}개 메모리 보관 중")
            if st.button("🗑️ Series JSON 로드 데이터 비우기", key="rw_clear_loaded"):
                st.session_state.pop("_rw_loaded_beats", None)
                st.session_state.pop("_rw_loaded_plans", None)
                st.rerun()

    # ── ① 기획서 (Creator Engine 산출물) ──
    st.markdown(
        '<div class="section-header">📋 기획서 <span class="en">CREATOR ENGINE OUTPUT</span></div>',
        unsafe_allow_html=True,
    )
    st.caption("Creator Engine에서 완성한 기획서를 붙여넣으세요. 없는 항목은 비워도 됩니다.")

    with st.expander("📝 기획서 붙여넣기", expanded=True):
        rw_logline = st.text_area("로그라인", height=80, placeholder="Logline Pack", key="rw_logline")
        rw_gns = st.text_area("GNS + 서사동력", height=100, placeholder="Goal/Need/Strategy + narrative_drive", key="rw_gns")
        rw_characters = st.text_area("캐릭터 + 바이블", height=180, placeholder="characters + extended_characters + Bible (tactics·secret·speech_pattern)", key="rw_characters")
        col_rw_a, col_rw_b = st.columns(2)
        with col_rw_a:
            rw_world = st.text_area("세계관", height=100, placeholder="World Building", key="rw_world")
        with col_rw_b:
            rw_tone = st.text_area("톤 문서", height=100, placeholder="Tone Document", key="rw_tone")

    # ── ② 기존 대본 ──
    st.markdown(
        '<div class="section-header">📖 기존 대본 <span class="en">EXISTING SCRIPT</span></div>',
        unsafe_allow_html=True,
    )
    st.caption("리라이트할 에피소드의 대본을 붙여넣으세요. 시즌 아크와 씬 플랜도 있으면 결과가 훨씬 좋습니다.")

    rw_genre = st.selectbox("장르", list(GENRE_RULES.keys()), key="rw_genre")

    col_ep, col_bt = st.columns(2)
    with col_ep:
        st.session_state.setdefault("rw_ep_num", 1)
        rw_ep_num = st.number_input(
            "에피소드 번호", min_value=1, max_value=12,
            key="rw_ep_num",
            on_change=_rw_sync_beat_from_loaded,  # ★ v2.0.5
        )
    with col_bt:
        rw_beat_num = st.selectbox(
            "비트 번호",
            list(range(8)),
            format_func=lambda x: f"Beat {x} — {EPISODE_BEATS[x]['name']}" if x < len(EPISODE_BEATS) else f"Beat {x}",
            key="rw_beat_num",
            on_change=_rw_sync_beat_from_loaded,  # ★ v2.0.5
        )

    rw_beat_text = st.text_area(
        f"EP{rw_ep_num} Beat {rw_beat_num} 기존 대본",
        height=300,
        placeholder="리라이트할 비트의 시나리오 텍스트를 여기에 붙여넣으세요.",
        key="rw_beat_text",
    )

    with st.expander("📑 시즌 아크 · 씬 플랜 (선택사항 — 있으면 품질 ↑)", expanded=False):
        rw_season_arc = st.text_area("시즌 아크", height=150, placeholder="시즌 아크 설계 텍스트 (없으면 비워두세요)", key="rw_season_arc")
        rw_episode_plan = st.text_area(f"EP{rw_ep_num} 씬 플랜", height=150, placeholder="해당 에피소드의 씬 플랜 (없으면 비워두세요)", key="rw_ep_plan")

    # ── ③ 후속 에피소드 (연속성 유지용) ──
    with st.expander("📌 후속 에피소드 텍스트 (연속성 유지용 — 선택사항)", expanded=False):
        st.caption(
            "리라이트 대상 이후의 에피소드(EP5~8 등) 텍스트를 붙여넣으면 "
            "리라이트가 후속 에피소드와 모순되지 않도록 검증합니다."
        )
        rw_subsequent = st.text_area(
            "후속 에피소드 텍스트",
            height=200,
            placeholder="EP5~EP8의 대본 텍스트 (요약도 가능)",
            key="rw_subsequent",
        )

    # ── ④ LOCKED / OPEN ──
    with st.expander("🔒 설정 잠금 (LOCKED / OPEN)", expanded=False):
        col_rl, col_ro = st.columns(2)
        with col_rl:
            rw_locked = st.text_area(
                "🔒 LOCKED (변경 불가)", height=120,
                placeholder="서재중: 29세, 묘적사 현장 요원\n기획의도: 20대 취업난이 입사 동기에 반영",
                key="rw_locked",
            )
        with col_ro:
            rw_open = st.text_area(
                "🟢 OPEN (창작 가능)", height=120,
                placeholder="캐릭터 외형, 말투 디테일\n장면별 연출",
                key="rw_open",
            )

    rw_locked_items = [l.strip() for l in rw_locked.strip().split("\n") if l.strip()] if rw_locked.strip() else []
    rw_open_items = [l.strip() for l in rw_open.strip().split("\n") if l.strip()] if rw_open.strip() else []
    rw_locked_block = build_locked_block(rw_locked_items, rw_open_items) if rw_locked_items or rw_open_items else ""

    # ── ⑤ 리라이트 노트 ──
    st.markdown(
        '<div class="section-header">📝 리라이트 노트 <span class="en">REWRITE NOTES</span></div>',
        unsafe_allow_html=True,
    )

    producer_notes = st.text_area(
        "🎬 프로듀서 / 넷플릭스 노트",
        height=150,
        placeholder=(
            "예:\n"
            "- EP1 오프닝이 너무 느리다. 3분 안에 사건으로.\n"
            "- 재중의 감정 변화가 급격하다. EP2까지 서서히.\n"
            "- B-Story 대선 카운트다운이 더 눈에 보여야."
        ),
        key="rw_producer",
    )

    col_n1, col_n2 = st.columns(2)
    with col_n1:
        casting_notes = st.text_area(
            "🎭 캐스팅 노트", height=120,
            placeholder="서재중 역: ○○○ 확정. 절제된 연기.\n최이호 역: △△△. 카리스마+유머.",
            key="rw_casting",
        )
    with col_n2:
        monitoring_notes = st.text_area(
            "📋 모니터링 / 내부 리뷰", height=120,
            placeholder="S#5 심문 씬이 가장 강하다. 유지.\nS#8~10 늘어진다. 압축 필요.",
            key="rw_monitoring",
        )

    # ── ⑥ 실행 ──
    has_text = bool(rw_beat_text.strip())
    has_notes = any([producer_notes.strip(), casting_notes.strip(), monitoring_notes.strip()])

    if not has_text:
        st.info("리라이트할 기존 대본 텍스트를 붙여넣으세요.")
    elif not has_notes:
        st.info("프로듀서 노트, 캐스팅 노트, 모니터링 의견 중 하나 이상을 입력하세요.")

    if st.button(
        f"🔄 EP{rw_ep_num} Beat {rw_beat_num} 구조적 리라이트 실행",
        type="primary",
        use_container_width=True,
        disabled=not (has_text and has_notes),
    ):
        prompt = build_structural_rewrite_prompt(
            beat_text=rw_beat_text,
            ep_num=rw_ep_num,
            beat_num=rw_beat_num,
            season_arc=rw_season_arc,
            episode_plan=rw_episode_plan,
            genre=rw_genre,
            character_bible=rw_characters,
            producer_notes=producer_notes,
            casting_notes=casting_notes,
            monitoring_notes=monitoring_notes,
            subsequent_eps_context=rw_subsequent,
            locked_block=rw_locked_block,
        )
        with st.spinner(f"EP{rw_ep_num} Beat {rw_beat_num}을 구조적으로 리라이트하고 있습니다..."):
            result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_BEAT)

        # 결과를 세션에 저장 (다운로드용)
        if "rw_results" not in st.session_state:
            st.session_state["rw_results"] = {}
        st.session_state["rw_results"][f"EP{rw_ep_num}_Beat{rw_beat_num}"] = result
        st.rerun()

    # ── 리라이트 결과 표시 + 다운로드 ──
    rw_results = st.session_state.get("rw_results", {})
    if rw_results:
        st.markdown("---")
        st.markdown(
            '<div class="section-header">📥 리라이트 결과 <span class="en">REWRITE RESULTS</span></div>',
            unsafe_allow_html=True,
        )
        timestamp = datetime.now().strftime("%Y%m%d")
        for label, text in rw_results.items():
            with st.expander(f"✅ {label}", expanded=False):
                st.markdown(text[:5000] + ("..." if len(text) > 5000 else ""))
            build_txt_download(text, f"{label}_rewrite_{timestamp}.txt")

        if st.button("🗑️ 리라이트 결과 초기화"):
            st.session_state["rw_results"] = {}
            st.rerun()

    st.markdown("---")
    st.caption(f"© 2026 BLUE JEANS PICTURES · Series Engine {ENGINE_VERSION}")
    st.stop()


# ══════════════════════════════════════════════
# ✍️ 집필 모드 (아래부터 기존 파이프라인)
# ══════════════════════════════════════════════

# STEP 1: 자료 입력
# ══════════════════════════════════════════════

st.markdown(
    '<div class="section-header">📋 자료 입력 <span class="en">STEP 1 · INPUT</span></div>',
    unsafe_allow_html=True,
)

col_s1, col_s2, col_s3 = st.columns(3)
with col_s1:
    st.session_state["num_episodes"] = st.selectbox(
        "🎞️ 에피소드 수", [6, 8],
        index=1 if st.session_state["num_episodes"] == 8 else 0,
    )
with col_s2:
    dur_options = [40, 50, 60]
    st.session_state["duration"] = st.selectbox(
        "⏱️ 에피소드 분량 (분)", dur_options,
        index=dur_options.index(st.session_state["duration"]),
    )
with col_s3:
    genre_options = list(GENRE_RULES.keys())
    st.session_state["genre"] = st.selectbox(
        "🎬 장르", genre_options,
        index=genre_options.index(st.session_state["genre"]) if st.session_state["genre"] in genre_options else 0,
    )

# LOCKED / OPEN 설정
with st.expander("🔒 설정 잠금 (LOCKED / OPEN)", expanded=False):
    st.caption(
        "LOCKED = 파이프라인 전 과정에서 절대 변경 불가. "
        "OPEN = AI가 자유롭게 창작 가능. 한 줄에 하나씩 입력하세요."
    )
    col_lock, col_open = st.columns(2)
    with col_lock:
        # ★ v1.8 패치: value= 제거 → setdefault로 초기값 주입 (안티패턴 회피)
        st.session_state.setdefault(
            "locked_input", "\n".join(st.session_state["locked_items"])
        )
        locked_input = st.text_area(
            "🔒 LOCKED (변경 불가)",
            height=150,
            placeholder=(
                "한 줄에 하나씩:\n"
                "서재중: 29세, 묘적사 현장 요원\n"
                "기획의도: 20대 취업난이 재중의 입사 동기에 반영\n"
                "B-Story: 대선 D-47부터 선거일까지 카운트다운"
            ),
            key="locked_input",
        )
    with col_open:
        st.session_state.setdefault(
            "open_input", "\n".join(st.session_state["open_items"])
        )
        open_input = st.text_area(
            "🟢 OPEN (창작 가능)",
            height=150,
            placeholder=(
                "한 줄에 하나씩:\n"
                "캐릭터 외형, 습관, 말투 디테일\n"
                "장면별 시각 연출과 감정 변화\n"
                "대사의 구체적 워딩"
            ),
            key="open_input",
        )
    st.session_state["locked_items"] = [l.strip() for l in locked_input.strip().split("\n") if l.strip()] if locked_input.strip() else []
    st.session_state["open_items"] = [l.strip() for l in open_input.strip().split("\n") if l.strip()] if open_input.strip() else []

INPUT_FIELDS = [
    ("logline",    "① 로그라인",           "Creator Engine의 Logline Pack (시리즈용, 시즌 질문 포함)"),
    ("intention",  "② 기획의도",           "Creator Engine의 KEY POINTS"),
    ("gns",        "③ GNS + 서사동력",     "Goal / Need / Strategy + narrative_drive (desire_origin·arc_direction·goal_need_gap)"),
    ("characters", "④ 캐릭터 + 바이블",    "characters(4인) + extended_characters + 바이블 (tactics·secret·speech_pattern·sample_lines)"),
    ("world",      "⑤ 세계관",             "World Building"),
    ("structure",  "⑥ 구조",               "Synopsis + Storyline + Beat Sheet (시즌 아크 기준)"),
    ("scenes",     "⑦ 장면 설계",          "에피소드별 핵심 장면"),
    ("treatment",  "⑧ 트리트먼트",         "에피소드별 분할 트리트먼트"),
    ("tone",       "⑨ 톤 문서",            "Tone Document (시리즈 톤 지속성 규칙 포함)"),
]

# ★ v1.8 — Creator Engine JSON 자동 로더
with st.expander("⚡ Creator Engine JSON 업로드 (자동 채우기)", expanded=False):
    st.markdown(
        '<div class="small-meta">Creator Engine에서 내려받은 .json 파일을 업로드하면 '
        '아래 9칸이 자동으로 채워집니다.</div>',
        unsafe_allow_html=True,
    )
    json_file = st.file_uploader("Creator Engine JSON 파일", type=["json"], key="creator_json_uploader")
    col_j1, col_j2 = st.columns([1, 2])
    with col_j1:
        load_json_btn = st.button("📂 JSON 적용", use_container_width=True,
                                   disabled=(json_file is None), key="json_load_btn")
    with col_j2:
        if json_file is not None:
            st.caption(f"선택됨: `{json_file.name}`")

    if load_json_btn and json_file is not None:
        try:
            raw = json_file.read().decode("utf-8")
            creator_data = json.loads(raw)
            loaded = extract_from_creator_json_series(creator_data)

            # ★ v1.8 패치: 위젯 key는 위에서 이미 그려졌을 수 있으므로 pending dict에 모음
            pending = st.session_state.get("_pending_widget_sync", {}) or {}

            # 9칸 입력에 주입
            for key, _, _ in INPUT_FIELDS:
                if loaded.get(key):
                    st.session_state["inputs"][key] = loaded[key]
                    pending[f"input_{key}"] = loaded[key]

            # ★ v2.0.5: 작품 제목 자동 주입 (Creator JSON의 project.title)
            loaded_title = (loaded.get("title", "") or "").strip()
            if loaded_title:
                st.session_state["inputs"]["title"] = loaded_title
                pending["input_title"] = loaded_title

            # v1.8: LOCKED 5종 확장 — Creator JSON에서 추출된 LOCKED 항목 자동 추가
            locked_ext = loaded.get("locked_5_extended", "")
            if locked_ext:
                existing = st.session_state.get("locked_items", [])
                new_items = [l.strip() for l in locked_ext.split("\n") if l.strip()]
                # 중복 방지
                for item in new_items:
                    if item not in existing:
                        existing.append(item)
                st.session_state["locked_items"] = existing
                # 위젯 key 동기화 (pending)
                pending["locked_input"] = "\n".join(existing)

            if pending:
                st.session_state["_pending_widget_sync"] = pending

            # 메타 정보
            meta = creator_data.get("_meta", {})
            ce_ver = meta.get("engine_version", "?")
            stage = meta.get("stage", "?")
            title = loaded.get("title", "(무제)")
            locked_count = len(locked_ext.split("\n")) if locked_ext else 0

            st.success(
                f"✅ Creator Engine {ce_ver} / {stage} 단계 로드 완료.\n\n"
                f"**프로젝트**: {title}\n\n"
                f"**로드된 필드**: {sum(1 for k, _, _ in INPUT_FIELDS if loaded.get(k))}개 / 9칸\n\n"
                f"**LOCKED 확장**: {locked_count}개 항목 추가" if locked_count else
                f"✅ Creator Engine {ce_ver} / {stage} 단계 로드 완료.\n\n"
                f"**프로젝트**: {title}\n\n"
                f"**로드된 필드**: {sum(1 for k, _, _ in INPUT_FIELDS if loaded.get(k))}개 / 9칸"
            )
            st.rerun()
        except json.JSONDecodeError as e:
            st.error(f"JSON 파싱 실패: {e}")
        except Exception as e:
            st.error(f"로드 중 오류: {e}")

with st.expander("📝 Creator Engine 결과 붙여넣기 (9칸)", expanded=not st.session_state["season_arc"]):
    # ★ v2.0.5 — 작품 제목 (백업 파일명·사이드바·PDF 푸터에 사용)
    _title_widget_key = "input_title"
    st.session_state.setdefault(_title_widget_key, st.session_state["inputs"].get("title", ""))
    st.text_input(
        "🎬 작품 제목",
        placeholder="예: 클레어의 서울 / 수비니어샵 / 왕게임 시즌 2",
        key=_title_widget_key,
        help="백업 JSON 파일명, 사이드바, PDF 푸터에 표시됩니다. 비워두면 장르 기반 파일명으로 폴백.",
    )
    st.session_state["inputs"]["title"] = st.session_state[_title_widget_key]

    for key, label, placeholder in INPUT_FIELDS:
        height = 200 if key == "characters" else 120
        # ★ v1.8 패치: value= 제거 → setdefault로 초기값 주입 (안티패턴 회피)
        widget_key = f"input_{key}"
        st.session_state.setdefault(widget_key, st.session_state["inputs"][key])
        st.text_area(
            label,
            height=height, placeholder=placeholder, key=widget_key,
        )
        # 위젯 값 → inputs 딕셔너리 동기화 (다른 영역에서 inputs[key]를 직접 읽으므로 유지)
        st.session_state["inputs"][key] = st.session_state[widget_key]


# ══════════════════════════════════════════════
# STEP 2: 시즌 아크
# ══════════════════════════════════════════════

st.markdown(
    '<div class="section-header">🏗️ 시즌 아크 설계 <span class="en">STEP 2 · SEASON ARC</span></div>',
    unsafe_allow_html=True,
)

ne = st.session_state["num_episodes"]
dur = st.session_state["duration"]
genre = st.session_state["genre"]

if st.session_state["season_arc"]:
    with st.expander("✅ 시즌 아크 (완료)", expanded=False):
        st.markdown(st.session_state["season_arc"])
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            build_txt_download(st.session_state["season_arc"], "season_arc.txt")
        with col_d2:
            build_docx_download(
                st.session_state["season_arc"], "season_arc.docx",
                title=f"시즌 아크 — {ne}부작 · {genre}",
            )
else:
    has_input = any(v.strip() for v in st.session_state["inputs"].values())

    st.markdown(
        f'<div class="callout">'
        f'<div class="cl">시즌 설정</div>'
        f'{ne}부작 · 에피소드당 {dur}분 · {genre}'
        f'</div>',
        unsafe_allow_html=True,
    )

    if st.button(f"🎬 {ne}부작 시즌 아크 설계 시작", disabled=not has_input, type="primary", use_container_width=True):
        prompt = build_season_arc_prompt(
            st.session_state["inputs"], ne, dur, genre,
            locked_block=_get_locked_block(),
            monitoring_feedback=st.session_state.get("monitoring_feedback", ""),
            showrunner_notes=st.session_state.get("showrunner_notes", ""),
        )
        with st.spinner("시즌 아크를 설계하고 있습니다..."):
            result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_ARC, model=MODEL_PLAN)
        st.session_state["season_arc"] = result
        st.rerun()

    # ★ v2.0.4 — 외부 피드백 + 쇼러너 노트 입력 (시즌 아크 미완성 상태에서만 노출)
    with st.expander("🗒️ 외부 모니터링 의견 / 쇼러너 노트 (선택, 시즌 아크 설계에 반영)", expanded=False):
        st.caption(
            "외부 검토자 모니터링 의견 또는 작가 메모를 입력하시면 시즌 아크 설계 시 자동으로 반영됩니다. "
            "비워두면 기존 자료만으로 설계합니다."
        )

        # 기존값 → setdefault로 안전하게 widget 초기화
        st.session_state.setdefault("monitoring_feedback", "")
        st.session_state.setdefault("showrunner_notes", "")

        st.text_area(
            "외부 모니터링 의견",
            height=200,
            placeholder=(
                "예시:\n"
                "검토자 총평: ...\n"
                "컨셉/주제: ...\n"
                "이야기 구조: ...\n"
                "캐릭터: ...\n"
                "대사·톤: ...\n"
                "기타: ..."
            ),
            key="monitoring_feedback",
            help="외부 검토자가 보낸 모니터링 의견을 그대로 붙여넣으세요. 6항목 형태로 구조화되어 있으면 더 정확합니다.",
        )

        st.text_area(
            "쇼러너 / 작가 노트 (최우선 반영)",
            height=150,
            placeholder=(
                "예시:\n"
                "- 검토자 의견 중 '진영 간 갈등 명확화'는 반영하되 '주인공 사진기억능력'은 수용하지 않음.\n"
                "- 정섬-덕문 과거사를 EP3 후반에 명시적으로 회수할 것.\n"
                "- 시즌 2 시드 30%만 남기고 시즌 1은 자체 완결성 강화."
            ),
            key="showrunner_notes",
            help="외부 의견 중 수용/거부 결정, 시즌 아크 설계 시 반드시 반영할 사항 등을 자유 형식으로 작성.",
        )


# ══════════════════════════════════════════════
# STEP 2.5: 시리즈 확장 (v1.7 — 캐릭터·사건 보강 + 핵심 요소 + Plant-Payoff)
# ══════════════════════════════════════════════

if st.session_state["season_arc"]:
    st.markdown(
        '<div class="section-header">🔧 시리즈 확장 <span class="en">STEP 2.5 · SERIES EXPANSION</span></div>',
        unsafe_allow_html=True,
    )

    # ── 2.5a: 캐릭터 확장 ──
    if not st.session_state.get("expanded_characters"):
        st.caption("① 시즌 아크를 분석하여 부족한 캐릭터를 생성합니다.")
        if st.button("👤 캐릭터 확장 분석", type="primary", use_container_width=True, key="char_expand"):
            prompt = build_character_expansion_prompt(
                st.session_state["inputs"],
                st.session_state["season_arc"],
                genre, ne,
                locked_block=_get_locked_block(),
            )
            with st.spinner("캐릭터 확장을 분석하고 있습니다..."):
                result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_ARC, model=MODEL_PLAN)
            st.session_state["expanded_characters"] = result
            st.rerun()
    else:
        with st.expander("👤 캐릭터 확장 (완료)", expanded=False):
            st.markdown(st.session_state["expanded_characters"])

    # ── 2.5b: 사건 보강 ──
    if st.session_state.get("expanded_characters") and not st.session_state.get("expanded_events"):
        st.caption("② 시즌 아크와 기존 트리트먼트를 대조하여 부족한 사건을 보강합니다.")
        if st.button("📋 사건 보강 분석", type="primary", use_container_width=True, key="event_expand"):
            prompt = build_event_expansion_prompt(
                st.session_state["inputs"],
                st.session_state["season_arc"],
                genre, ne,
                locked_block=_get_locked_block(),
            )
            with st.spinner("사건을 보강하고 있습니다..."):
                result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_ARC, model=MODEL_PLAN)
            st.session_state["expanded_events"] = result
            st.rerun()
    elif st.session_state.get("expanded_events"):
        with st.expander("📋 사건 보강 (완료)", expanded=False):
            st.markdown(st.session_state["expanded_events"])

    # ── 2.5c: 핵심 요소 추출 + Plant-Payoff 맵 ──
    if st.session_state.get("expanded_events") and not st.session_state["story_elements"]:
        st.caption("③ 맥거핀 · 비밀 · 전술 · 장소 · 모티프 · Plant-Payoff 맵을 추출합니다.")
        if st.button("🔍 핵심 요소 + Plant-Payoff 추출", type="primary", use_container_width=True, key="extract_elements"):
            prompt = build_extract_elements_prompt(
                st.session_state["inputs"],
                st.session_state["season_arc"],
                genre,
                locked_block=_get_locked_block(),
            )
            with st.spinner("핵심 요소를 추출하고 있습니다..."):
                result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_ARC, model=MODEL_PLAN)
            st.session_state["story_elements"] = result
            st.rerun()
    elif st.session_state["story_elements"]:
        with st.expander("🔍 핵심 요소 + Plant-Payoff (완료)", expanded=False):
            st.markdown(st.session_state["story_elements"])


# ══════════════════════════════════════════════
# STEP 3: 에피소드별 씬 플랜
# ══════════════════════════════════════════════

if st.session_state["season_arc"]:
    st.markdown(
        '<div class="section-header">📑 에피소드별 씬 플랜 <span class="en">STEP 3 · SCENE PLAN</span></div>',
        unsafe_allow_html=True,
    )

    ep_html = '<div class="ep-tabs">'
    for ep in range(1, ne + 1):
        if ep in st.session_state["episode_plans"]:
            ep_html += f'<div class="ep-tab done">EP{ep} ✓</div>'
        else:
            prev_done = (ep == 1) or ((ep - 1) in st.session_state["episode_plans"])
            cls = "current" if prev_done else "pending"
            ep_html += f'<div class="ep-tab {cls}">EP{ep}</div>'
    ep_html += '</div>'
    st.markdown(ep_html, unsafe_allow_html=True)

    for ep in range(1, ne + 1):
        if ep in st.session_state["episode_plans"]:
            with st.expander(f"✅ EP{ep} 씬 플랜", expanded=False):
                st.markdown(st.session_state["episode_plans"][ep])
                build_txt_download(st.session_state["episode_plans"][ep], f"EP{ep}_scene_plan.txt")

                # v1.7.1: 에피소드별 인물 추가 (게스트/카메오/기능적 조연)
                ep_char_key = f"ep_char_{ep}"
                if st.session_state.get("ep_characters", {}).get(ep):
                    st.markdown("---")
                    st.markdown(f"**👤 EP{ep} 추가 인물**")
                    st.markdown(st.session_state["ep_characters"][ep])
                
                if st.button(f"👤 EP{ep} 인물 추가/갱신", key=f"ep_char_btn_{ep}", use_container_width=True):
                    prompt = build_episode_character_prompt(
                        st.session_state["inputs"],
                        st.session_state["season_arc"],
                        st.session_state["episode_plans"][ep],
                        ep, genre,
                        existing_characters=st.session_state["inputs"].get("characters", ""),
                        expanded_characters=st.session_state.get("expanded_characters", ""),
                        locked_block=_get_locked_block(),
                    )
                    with st.spinner(f"EP{ep} 인물을 설계하고 있습니다..."):
                        result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_ARC, model=MODEL_PLAN)
                    if "ep_characters" not in st.session_state:
                        st.session_state["ep_characters"] = {}
                    st.session_state["ep_characters"][ep] = result
                    st.rerun()

    next_ep = None
    for ep in range(1, ne + 1):
        if ep not in st.session_state["episode_plans"]:
            next_ep = ep
            break

    if next_ep:
        prev_plan = st.session_state["episode_plans"].get(next_ep - 1, "")
        prev_last_scene = ""
        if next_ep > 1 and bk(next_ep - 1, 7) in st.session_state["episode_beats"]:
            full = st.session_state["episode_beats"][bk(next_ep - 1, 7)]
            prev_last_scene = full[-1500:] if len(full) > 1500 else full

        beats_ref = SEASON_BEATS_8 if ne == 8 else SEASON_BEATS_6
        season_beat_info = ""
        for b in beats_ref:
            if b["ep"] == next_ep:
                season_beat_info = f"{b['beat']} — {b['role']}"
                break

        st.markdown(
            f'<div class="callout">'
            f'<div class="cl">EP{next_ep}</div>'
            f'{season_beat_info}'
            f'</div>',
            unsafe_allow_html=True,
        )

        if st.button(
            f"📝 EP{next_ep} 씬 플랜 생성",
            key=f"gen_plan_{next_ep}",
            type="primary",
            use_container_width=True,
        ):
            prompt = build_episode_plan_prompt(
                st.session_state["inputs"],
                st.session_state["season_arc"],
                next_ep, ne, dur, genre,
                prev_episode_plan=prev_plan,
                prev_episode_last_scene=prev_last_scene,
                locked_block=_get_locked_block(),
            )
            with st.spinner(f"EP{next_ep} 씬 플랜을 작성하고 있습니다..."):
                result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_PLAN, model=MODEL_PLAN)
            st.session_state["episode_plans"][next_ep] = result
            st.rerun()


# ══════════════════════════════════════════════
# STEP 4: 비트별 집필
# ══════════════════════════════════════════════

if st.session_state["episode_plans"]:
    st.markdown(
        '<div class="section-header">✍️ 비트별 집필 <span class="en">STEP 4 · WRITING</span></div>',
        unsafe_allow_html=True,
    )

    # v1.7: 프로듀서 노트 (집필모드)
    with st.expander("🎬 프로듀서 노트 (집필 시 반영)", expanded=False):
        st.caption("여기에 입력한 노트가 모든 비트 집필에 자동 주입됩니다.")
        st.session_state["producer_notes_write"] = st.text_area(
            "프로듀서 노트",
            value=st.session_state.get("producer_notes_write", ""),
            height=100,
            placeholder="예: 대사를 더 날카롭게, B-Story 비중을 늘려라, EP3부터 긴장감 ↑",
            key="producer_notes_write_input",
        )

    available_eps = sorted(st.session_state["episode_plans"].keys())
    selected_ep = st.selectbox(
        "집필할 에피소드 선택",
        available_eps,
        format_func=lambda x: f"EP{x}",
        key="write_ep_select",
    )

    if selected_ep:
        next_beat = None
        beat_html = '<div class="beat-row">'
        for b in range(8):
            key = bk(selected_ep, b)
            if key in st.session_state["episode_beats"]:
                beat_html += f'<div class="beat-dot done">✓</div>'
            elif next_beat is None:
                beat_html += f'<div class="beat-dot current">{b}</div>'
                next_beat = b
            else:
                beat_html += f'<div class="beat-dot pending">{b}</div>'
        beat_html += '</div>'
        st.markdown(beat_html, unsafe_allow_html=True)

        st.caption(" · ".join([f"B{b['beat']}={b['name']}" for b in EPISODE_BEATS]))

        for b in range(8):
            key = bk(selected_ep, b)
            if key in st.session_state["episode_beats"]:
                beat_name = EPISODE_BEATS[b]["name"] if b < len(EPISODE_BEATS) else f"Beat {b}"
                with st.expander(f"✅ Beat {b} — {beat_name}", expanded=False):
                    st.markdown(st.session_state["episode_beats"][key])

        if next_beat is not None:
            beat_info = EPISODE_BEATS[next_beat] if next_beat < len(EPISODE_BEATS) else EPISODE_BEATS[-1]

            st.markdown(
                f'<div class="callout" style="border-left-color:var(--y)">'
                f'<div class="cl">EP{selected_ep} · Beat {next_beat}</div>'
                f'{beat_info["name"]} — {beat_info["minutes"]} — {beat_info["role"]}'
                f'</div>',
                unsafe_allow_html=True,
            )

            prev_beat_text = ""
            if next_beat > 0:
                prev_beat_text = st.session_state["episode_beats"].get(bk(selected_ep, next_beat - 1), "")
            elif selected_ep > 1:
                prev_beat_text = st.session_state["episode_beats"].get(bk(selected_ep - 1, 7), "")

            char_bible = st.session_state["inputs"].get("characters", "")
            # v1.7.1: 에피소드별 추가 인물 정보를 캐릭터 바이블에 합산
            ep_extra_chars = st.session_state.get("ep_characters", {}).get(selected_ep, "")
            if ep_extra_chars:
                char_bible += f"\n\n[EP{selected_ep} 추가 인물 — 게스트/카메오/기능적 조연]\n{ep_extra_chars}"
            # 시즌 확장 캐릭터도 합산
            expanded_chars = st.session_state.get("expanded_characters", "")
            if expanded_chars:
                char_bible += f"\n\n[시리즈 확장 캐릭터]\n{expanded_chars[:2000]}"

            # v1.8: Profession Pack — 캐릭터 직업 디테일 자동 주입
            profession_block = ""
            try:
                # 캐릭터 바이블에서 직업 정보 감지
                prof_text = st.session_state["inputs"].get("characters", "")
                if prof_text:
                    profession_block = build_multi_profession_block([
                        {"name": "", "occupation": prof_text}
                    ])
                    if profession_block:
                        char_bible += f"\n\n{profession_block[:3000]}"
            except Exception:
                pass

            # v1.8: Period Pack — 시대 디테일 자동 주입
            period_block = ""
            try:
                locked_text = "\n".join(st.session_state.get("locked_items", []))
                world_text = st.session_state["inputs"].get("world", "")
                combined_text = f"{locked_text}\n{world_text}"
                if combined_text.strip():
                    period_block = build_period_block_auto(combined_text)
                    if period_block:
                        char_bible += f"\n\n{period_block[:2000]}"
            except Exception:
                pass

            if st.button(
                f"🖊️ EP{selected_ep} Beat {next_beat} [{beat_info['name']}] 집필",
                key=f"write_{selected_ep}_{next_beat}",
                type="primary",
                use_container_width=True,
            ):
                # v1.7: 비트 구조 유형 추적
                prev_struct = st.session_state.get("beat_structure_types", {}).get(
                    bk(selected_ep, next_beat - 1), ""
                ) if next_beat > 0 else ""
                # v1.8: 적응형 컨텍스트 관리 — EP 진행에 따라 정보량 자동 조절
                ep_summary = ""
                if selected_ep > 1:
                    # 적응형: EP 진행에 따라 이전 에피소드 정보량 조절
                    ep_summaries = st.session_state.get("episode_summaries", {})
                    # 아직 요약이 없는 에피소드는 생성
                    for prev_ep in range(1, selected_ep):
                        if prev_ep not in ep_summaries:
                            s = summarize_episode_context(st.session_state["episode_beats"], prev_ep)
                            if s:
                                if "episode_summaries" not in st.session_state:
                                    st.session_state["episode_summaries"] = {}
                                st.session_state["episode_summaries"][prev_ep] = s
                    ep_summary = build_adaptive_context(
                        st.session_state["episode_beats"],
                        selected_ep,
                        st.session_state.get("episode_summaries", {}),
                    )

                # v1.8: 비트 간 중복 방지 — 이전 비트 정보를 inputs에 임시 추가
                inputs_with_beats = dict(st.session_state["inputs"])
                inputs_with_beats["_episode_beats"] = st.session_state.get("episode_beats", {})

                # ★ v2.0 — 시즌 표현 누적 DB (없으면 빈 dict)
                if "season_expression_db" not in st.session_state:
                    st.session_state["season_expression_db"] = {}
                _season_expr_db = st.session_state["season_expression_db"]

                prompt = build_write_episode_beat_prompt(
                    inputs_with_beats,
                    st.session_state["season_arc"],
                    st.session_state["episode_plans"][selected_ep],
                    selected_ep, next_beat, ne, dur, genre,
                    prev_beat_text=prev_beat_text,
                    character_bible=char_bible,
                    story_elements=st.session_state.get("story_elements", ""),
                    locked_block=_get_locked_block(),
                    producer_notes=st.session_state.get("producer_notes_write", ""),
                    prev_beat_structure_type=prev_struct,
                    episode_context_summary=ep_summary,
                    season_expression_db=_season_expr_db,  # v2.0
                    # ★ v2.0.8 — EP 첫 비트(Beat 0)인 경우 씬 번호 S#1부터 리셋
                    is_first_beat_of_episode=(next_beat == 0),
                )
                with st.spinner(f"EP{selected_ep} Beat {next_beat}을 집필하고 있습니다..."):
                    result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_BEAT)
                st.session_state["episode_beats"][bk(selected_ep, next_beat)] = result

                # ★ v2.0 — 시즌 표현 누적 DB 갱신
                update_season_expression_db(
                    st.session_state["season_expression_db"],
                    result, selected_ep, next_beat,
                )

                # v1.7: 비트 구조 유형 자동 추출
                if "beat_structure_types" not in st.session_state:
                    st.session_state["beat_structure_types"] = {}
                for code in ["INV", "CON", "REV", "EMO", "ACT", "SIL"]:
                    if f"[{code}]" in result:
                        st.session_state["beat_structure_types"][bk(selected_ep, next_beat)] = code
                        break
                st.rerun()
        else:
            st.markdown(
                f'<div class="callout" style="border-left-color:var(--g)">'
                f'<div class="cl">🎉 완료</div>'
                f'EP{selected_ep} 전체 집필 완료!'
                f'</div>',
                unsafe_allow_html=True,
            )

        # 다시 쓰기
        st.markdown("---")
        last_done = None
        for b in range(7, -1, -1):
            if bk(selected_ep, b) in st.session_state["episode_beats"]:
                last_done = b
                break

        if last_done is not None:
            st.markdown('<div class="cl">🔄 마지막 비트 다시 쓰기</div>', unsafe_allow_html=True)
            rewrite_instruction = st.text_area(
                f"EP{selected_ep} Beat {last_done} 수정 지시",
                placeholder="예: 대사를 더 날카롭게, 클리프행어를 Betrayal로 변경, S#15 삭제...",
                height=80,
                key="rewrite_inst",
            )
            if st.button(
                f"🔄 Beat {last_done} 다시 쓰기",
                key="rewrite_btn",
                disabled=not rewrite_instruction.strip() if isinstance(rewrite_instruction, str) else True,
            ):
                original = st.session_state["episode_beats"][bk(selected_ep, last_done)]
                prompt = build_rewrite_prompt(original, rewrite_instruction, genre, character_bible=char_bible, locked_block=_get_locked_block())
                with st.spinner(f"Beat {last_done}을 다시 쓰고 있습니다..."):
                    result = stream_response(SYSTEM_PROMPT, prompt, MAX_TOKENS_REWRITE)
                st.session_state["episode_beats"][bk(selected_ep, last_done)] = result
                st.rerun()


# ══════════════════════════════════════════════
# STEP 5: 다운로드
# ══════════════════════════════════════════════

has_any_beats = len(st.session_state["episode_beats"]) > 0

if has_any_beats:
    st.markdown(
        '<div class="section-header">💾 다운로드 <span class="en">STEP 5 · EXPORT</span></div>',
        unsafe_allow_html=True,
    )

    # ★ v2.0.1 — 파일명 규칙: 각본_제목_v1.0_날짜_시간.확장자
    # 제목은 Creator Engine 입력의 title 또는 logline 첫 어구에서 추출.
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    # 작품 제목 추출 (안전 가드)
    def _extract_title():
        """작품 제목을 안전하게 추출. 우선순위:
        1. inputs['title'] 명시값
        2. logline 앞부분 (가제 표기 제거 후 첫 어구)
        3. fallback: '제목미정'
        """
        _inp = st.session_state.get("inputs", {}) or {}
        # 1. title 명시 입력
        _title = (_inp.get("title") or "").strip()
        if _title:
            return _sanitize_filename(_title)
        # 2. logline에서 작품명 추출 시도
        _logline = (_inp.get("logline") or "").strip()
        if _logline:
            import re as _re_t
            # "제목:", "가제:", "「제목」", "<제목>", "[제목]" 패턴
            for _pat in [
                r'「([^」]+)」', r'<([^>]+)>', r'\[([^\]]+)\]',
                r'^(?:제목|가제)\s*[:：]\s*([^\n]+?)(?:\(|$|,|\n)',
            ]:
                _m = _re_t.search(_pat, _logline)
                if _m:
                    return _sanitize_filename(_m.group(1).strip())
        # 3. fallback
        return "제목미정"

    def _sanitize_filename(s: str) -> str:
        """파일명에 쓸 수 없는 문자 제거 + 공백 정리."""
        import re as _re_s
        s = _re_s.sub(r'[\\/:*?"<>|]', '', s)
        s = _re_s.sub(r'\s+', '_', s.strip())
        # 너무 길면 자르기
        if len(s) > 30:
            s = s[:30]
        return s or "제목미정"

    _work_title = _extract_title()
    _version_tag = "v1.0"  # 추후 버전 입력 UI 추가 시 확장

    # ★ v2.0.3 — DOCX 출력 모드 선택
    _output_mode = st.radio(
        "DOCX 출력 모드",
        options=["최종 모드 (제작·연출 전달용)", "집필 모드 (작가 검토용)"],
        index=0,
        horizontal=True,
        key="docx_output_mode",
        help=(
            "최종 모드: 비트 헤더(Beat 0~7) 제거 — 시나리오 표준 양식. 제작·투자·연출 전달용.\n"
            "집필 모드: 비트 헤더 포함 — 작가가 구조 단위로 검토할 때 사용."
        ),
    )
    _include_beats = "집필 모드" in _output_mode

    st.markdown('<div class="cl">에피소드별</div>', unsafe_allow_html=True)
    dl_cols = st.columns(min(ne, 4))
    for idx, ep in enumerate(range(1, ne + 1)):
        ep_text = get_all_episode_text(ep)
        if ep_text:
            with dl_cols[idx % min(ne, 4)]:
                _ep_fname_base = f"각본_{_work_title}_EP{ep}_{_version_tag}_{timestamp}"
                build_txt_download(ep_text, f"{_ep_fname_base}.txt")
                build_docx_download(
                    ep_text, f"{_ep_fname_base}.docx", title=f"EPISODE {ep}",
                    include_beat_headers=_include_beats,
                )

    full_text = get_full_season_text()
    if full_text:
        st.markdown('<div class="cl" style="margin-top:1rem">시즌 전체</div>', unsafe_allow_html=True)
        col_f1, col_f2 = st.columns(2)
        _season_fname_base = f"각본_{_work_title}_{_version_tag}_{timestamp}"
        with col_f1:
            build_txt_download(full_text, f"{_season_fname_base}.txt")
        with col_f2:
            build_docx_download(
                full_text, f"{_season_fname_base}.docx",
                title=f"시즌 1 — {ne}부작 · {genre}",
                include_beat_headers=_include_beats,
            )


# ══════════════════════════════════════════════
# ★ v1.8 — 프로젝트 세션 백업 (중단 시 복구)
# ══════════════════════════════════════════════

# 백업 대상 키
_BACKUP_KEYS = [
    "inputs", "num_episodes", "duration", "genre",
    "season_arc", "story_elements",
    "expanded_characters", "expanded_events",
    "episode_plans", "episode_beats",
    "beat_structure_types", "episode_summaries",
    "ep_characters",
    "locked_items", "open_items",
    "producer_notes_write",
    "season_expression_db",  # v2.0
    "monitoring_feedback",   # v2.0.4
    "showrunner_notes",      # v2.0.4
]


def _export_session_backup() -> bytes:
    """현재 세션 상태를 JSON bytes로 직렬화."""
    _inputs_snap = st.session_state.get("inputs", {}) or {}
    payload = {
        "_meta": {
            "engine": "Series Engine",
            "engine_version": ENGINE_VERSION,
            "saved_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
            "project_title": (_inputs_snap.get("title") or "").strip(),  # ★ v2.0.5
            "genre": st.session_state.get("genre", ""),
            "num_episodes": st.session_state.get("num_episodes", 8),
            "season_arc_done": bool(st.session_state.get("season_arc")),
            "episodes_planned": len(st.session_state.get("episode_plans", {})),
            "beats_written": len(st.session_state.get("episode_beats", {})),
        },
        "session": {k: st.session_state.get(k) for k in _BACKUP_KEYS},
    }
    # 딕셔너리 키를 str로 변환 (JSON 호환)
    for dict_key in ["episode_plans", "episode_beats", "beat_structure_types", "episode_summaries", "ep_characters"]:
        val = payload["session"].get(dict_key, {})
        if isinstance(val, dict):
            payload["session"][dict_key] = {str(k): v for k, v in val.items()}
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _import_session_backup(raw_bytes: bytes) -> dict:
    """JSON bytes를 받아 세션에 복원."""
    data = json.loads(raw_bytes.decode("utf-8"))
    session_data = data.get("session", {})
    meta = data.get("_meta", {})

    # ★ v2.0.6 패치: 위젯 key로 사용되는 신규 백업 키는 직접 대입 금지.
    # STEP 2 영역에서 이미 instantiate된 text_area들이라
    # 직접 st.session_state[k] = v 하면 StreamlitAPIException 발생.
    _WIDGET_KEYS_IN_BACKUP = {"monitoring_feedback", "showrunner_notes"}

    for k in _BACKUP_KEYS:
        if k in session_data:
            v = session_data[k]
            # 딕셔너리 키를 int로 복원 (에피소드 번호)
            if k in ["episode_plans", "ep_characters"] and isinstance(v, dict):
                v = {int(kk): vv for kk, vv in v.items()}
            # 위젯 key 항목은 pending으로 우회 (아래에서 처리)
            if k in _WIDGET_KEYS_IN_BACKUP:
                continue
            st.session_state[k] = v

    # ★ v1.8 패치: 위젯 key는 이미 그려진 상태라 직접 수정 불가.
    # _pending_widget_sync에 모아두고 rerun 후 페이지 최상단에서 적용.
    pending = {}
    # 9칸 입력 영역
    inputs_restored = session_data.get("inputs", {})
    if isinstance(inputs_restored, dict):
        for ik, iv in inputs_restored.items():
            if iv:
                pending[f"input_{ik}"] = iv
    # LOCKED / OPEN 입력 영역
    locked_restored = session_data.get("locked_items", [])
    if isinstance(locked_restored, list) and locked_restored:
        pending["locked_input"] = "\n".join(locked_restored)
    open_restored = session_data.get("open_items", [])
    if isinstance(open_restored, list) and open_restored:
        pending["open_input"] = "\n".join(open_restored)

    # ★ v2.0.6 — v2.0.4 신규 위젯도 pending으로 우회 (text_area는 빈 문자열도 안전)
    for _wk in _WIDGET_KEYS_IN_BACKUP:
        if _wk in session_data:
            pending[_wk] = session_data[_wk] or ""

    if pending:
        st.session_state["_pending_widget_sync"] = pending

    return meta


def _sanitize_for_filename(text: str, max_len: int = 30) -> str:
    """★ v2.0.5 — 백업 파일명용 sanitize (모듈 스코프).
    OS 호환: 위험 문자 제거, 공백→언더스코어, 길이 제한.
    """
    if not text:
        return ""
    for bad in ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '\n', '\r', '\t']:
        text = text.replace(bad, "")
    text = "_".join(text.split())
    text = text.strip("._")
    return text[:max_len]


def _make_backup_filename() -> str:
    # ★ v2.0.5 — 작품 제목 우선 패턴
    inputs = st.session_state.get("inputs", {}) or {}
    title = _sanitize_for_filename((inputs.get("title") or "").strip())
    genre = _sanitize_for_filename(st.session_state.get("genre", ""), max_len=20)
    ne = st.session_state.get("num_episodes", 8)
    eps = len(st.session_state.get("episode_plans", {}))
    beats = len(st.session_state.get("episode_beats", {}))
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    if title:
        return f"SeriesEngine_{title}_{genre}_{ne}ep_{eps}plans_{beats}beats_{ts}.json"
    return f"SeriesEngine_{genre}_{ne}ep_{eps}plans_{beats}beats_{ts}.json"


st.markdown("---")
with st.expander("💾 프로젝트 세션 백업 (중단 시 복구용)", expanded=False):
    st.markdown(
        '<div class="small-meta">현재 작업 중인 모든 입력칸·시즌 아크·씬 플랜·집필된 비트를 JSON으로 저장하거나 불러옵니다. '
        '비트 집필 도중 멈추거나 다음 날 이어서 작업할 때 사용하세요.</div>',
        unsafe_allow_html=True,
    )

    col_b1, col_b2 = st.columns(2)

    # ── 저장 ──
    with col_b1:
        st.markdown("**📥 백업 저장**")
        _bk_bytes = _export_session_backup()
        _bk_fname = _make_backup_filename()
        st.download_button(
            label=f"💾 JSON 다운로드",
            data=_bk_bytes,
            file_name=_bk_fname,
            mime="application/json",
            use_container_width=True,
            key="backup_download_btn",
        )
        st.caption(f"`{_bk_fname}`")

    # ── 불러오기 ──
    with col_b2:
        st.markdown("**📤 백업 불러오기**")
        backup_file = st.file_uploader(
            "백업 JSON 파일", type=["json"],
            key="backup_uploader",
            label_visibility="collapsed",
        )
        load_backup_btn = st.button(
            "📂 백업 적용 (현재 작업 덮어쓰기)",
            use_container_width=True,
            disabled=(backup_file is None),
            key="backup_load_btn",
        )

        if load_backup_btn and backup_file is not None:
            try:
                meta = _import_session_backup(backup_file.read())
                saved_ver = meta.get("engine_version", "?")
                saved_at = meta.get("saved_at", "?")
                eps_planned = meta.get("episodes_planned", "?")
                beats_written = meta.get("beats_written", "?")
                proj_title = (meta.get("project_title") or "").strip() or "(제목 없음)"  # ★ v2.0.5

                st.success(
                    f"✅ 백업 복원 완료\n\n"
                    f"**작품**: {proj_title}\n\n"
                    f"**저장 시각**: {saved_at}\n\n"
                    f"**엔진 버전**: {saved_ver}\n\n"
                    f"**씬 플랜**: {eps_planned}개 EP\n\n"
                    f"**집필 비트**: {beats_written}개"
                )
                st.rerun()
            except json.JSONDecodeError as e:
                st.error(f"JSON 파싱 실패: {e}")
            except Exception as e:
                st.error(f"복원 중 오류: {e}")


# ══════════════════════════════════════════════
# 전체 초기화
# ══════════════════════════════════════════════

st.markdown("---")
with st.expander("⚠️ 전체 초기화", expanded=False):
    st.warning("모든 진행 사항이 삭제됩니다.")
    if st.button("🗑️ 전체 초기화 실행"):
        for k, v in DEFAULTS.items():
            if isinstance(v, dict):
                st.session_state[k] = dict(v)
            else:
                st.session_state[k] = v
        st.rerun()

st.markdown("---")
st.caption(f"© 2026 BLUE JEANS PICTURES · Series Engine {ENGINE_VERSION}")
